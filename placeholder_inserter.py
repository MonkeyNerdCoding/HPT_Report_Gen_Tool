from __future__ import annotations

from collections.abc import Callable
import shutil
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from models import MappingRule, OperationCancelled


CancelCheck = Callable[[], bool]


@dataclass
class PlaceholderInsertReport:
    inserted: list[str] = field(default_factory=list)
    already_present: list[str] = field(default_factory=list)
    missing_anchors: list[str] = field(default_factory=list)
    backup_path: Path | None = None


def insert_mapping_placeholders(
    word_file: str | Path,
    rules: list[MappingRule],
    create_backup: bool = True,
    cancel_check: CancelCheck | None = None,
) -> PlaceholderInsertReport:
    path = Path(word_file)
    doc = Document(str(path))
    report = PlaceholderInsertReport()

    _raise_if_cancelled(cancel_check)

    if create_backup:
        backup = path.with_name(f"{path.stem}.placeholder_backup{path.suffix}")
        if not backup.exists():
            shutil.copy2(path, backup)
        report.backup_path = backup

    anchors = _build_anchor_index(doc)
    insertion_cursors: dict[int, Paragraph] = {}

    for rule in rules:
        _raise_if_cancelled(cancel_check)
        placeholder = rule.placeholder
        existing = _find_placeholder(doc, placeholder)
        if existing is not None:
            report.already_present.append(placeholder)
            continue

        anchor = _find_best_anchor(doc, anchors, rule)
        if anchor is None:
            report.missing_anchors.append(placeholder)
            continue

        insertion_anchor = insertion_cursors.get(id(anchor), anchor)
        insertion_cursors[id(anchor)] = _insert_after(insertion_anchor, placeholder)
        report.inserted.append(placeholder)

    if report.inserted:
        doc.save(str(path))

    return report


def _find_placeholder(doc: DocumentObject, placeholder: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            return paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        return paragraph
    return None


def _insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.style = "Normal"
    inserted.add_run(text)
    return inserted


def _build_anchor_index(doc: DocumentObject) -> dict[str, Paragraph]:
    anchors: dict[str, Paragraph] = {}
    for paragraph in doc.paragraphs:
        text = _normalize(paragraph.text)
        if text and text not in anchors:
            anchors[text] = paragraph
    return anchors


def _find_best_anchor(
    doc: DocumentObject,
    anchors: dict[str, Paragraph],
    rule: MappingRule,
) -> Paragraph | None:
    placeholder = rule.placeholder
    strategy = _PLACEHOLDER_STRATEGIES.get(placeholder, ())
    for candidate in strategy:
        if isinstance(candidate, tuple):
            anchor = _section_tail(doc, candidate[0], candidate[1])
            if anchor is not None:
                return anchor
            continue

        anchor = anchors.get(_normalize(candidate))
        if anchor is not None:
            return anchor

    for candidate in (rule.location, rule.section, rule.report_section):
        if not candidate:
            continue
        anchor = anchors.get(_normalize(candidate))
        if anchor is not None:
            return anchor

    if placeholder in _DATABASE_FALLBACK_PLACEHOLDERS:
        return anchors.get(_normalize("QUẢN LÝ DATABASE"))

    if placeholder in _MEMORY_FALLBACK_PLACEHOLDERS:
        return (
            anchors.get(_normalize("CẤU HÌNH VÙNG NHỚ CHO CSDL"))
            or anchors.get(_normalize("HIỆU SUẤT BỘ NHỚ"))
        )

    if placeholder in _ASH_FALLBACK_PLACEHOLDERS:
        return anchors.get(_normalize("ASH"))

    return None


def _raise_if_cancelled(cancel_check: CancelCheck | None) -> None:
    if cancel_check and cancel_check():
        raise OperationCancelled("Operation cancelled.")


def _section_tail(
    doc: DocumentObject,
    heading_text: str,
    next_heading_texts: tuple[str, ...],
) -> Paragraph | None:
    heading = _normalize(heading_text)
    next_headings = {_normalize(text) for text in next_heading_texts}
    in_section = False
    last_nonempty: Paragraph | None = None

    for paragraph in doc.paragraphs:
        text = _normalize(paragraph.text)
        if text == heading:
            in_section = True
            last_nonempty = paragraph
            continue
        if in_section and text in next_headings:
            return last_nonempty
        if in_section and text:
            last_nonempty = paragraph

    return last_nonempty if in_section else None


def _normalize(value: str) -> str:
    text = unicodedata.normalize("NFC", value or "")
    text = " ".join(text.strip().upper().split())
    return text


_DATABASE_FALLBACK_PLACEHOLDERS = {
    "<invalid_obj>",
    "<db_job>",
    "<sche_job>",
    "<table_stats>",
    "<index_stats>",
}

_MEMORY_FALLBACK_PLACEHOLDERS = {
    "<SGA_chart>",
    "<PGA_chart>",
    "<MEMORY_chart>",
}

_ASH_FALLBACK_PLACEHOLDERS = {
    "<aas_per_wait_class_for_instance_1>",
    "<aas_per_wait_class_for_instance_1_chart>",
    "<ash_top_timed_events_for_instance_1_for_9_days_of_history>",
    "<ash_top_timed_events_for_instance_1_for_9_days_of_history_chart>",
    "{{table_ash_top_sql_cluster_9_days}}",
    "{{table_ash_cpu_per_source}}",
    "<ash_cpu_per_source_chart>",
    "<ash_top_sql_for_instance_1_for_9_days_of_history_chart>",
}

_PLACEHOLDER_STRATEGIES: dict[str, tuple[str | tuple[str, tuple[str, ...]], ...]] = {
    "<tbs_usage>": ("MỨC ĐỘ SỬ DỤNG TABLESPACES",),
    "<data_file>": ("DATA FILE",),
    "<no_index_table>": (
        "CÁC BẢNG KHÔNG CÓ INDEX",
        "CÁC TABLES KHÔNG CÓ INDEX",
        "CÁC TABLE KHÔNG CÓ INDEX",
    ),
    "<no_pk_table>": (
        "CÁC BẢNG KHÔNG CÓ KHÓA CHÍNH",
        "CÁC BẢNG KHÔNG CÓ KHOÁ CHÍNH",
        "CÁC TABLE KHÔNG CÓ KHÓA CHÍNH",
        "CÁC TABLE KHÔNG CÓ KHOÁ CHÍNH",
        "CÁC TABLES KHÔNG CÓ KHÓA CHÍNH",
        "CÁC TABLES KHÔNG CÓ KHOÁ CHÍNH",
    ),
    "<invalid_obj>": ("INVALID OBJECT",),
    "<db_job>": ("Jobs",),
    "<sche_job>": ("Schedule Jobs", "Scheduler jobs"),
    "<table_stats>": ("Tables with Stale Stats", "TABLE VÀ INDEX CÓ STATISTICS CŨ"),
    "<index_stats>": ("Indexes with Stale Stats", "TABLE VÀ INDEX CÓ STATISTICS CŨ"),
    "<buffer_chart>": ("BUFFER CACHE HIT",),
    "<library_chart>": ("LIBRARY CACHE HIT",),
    "<SGA_chart>": ("SGA STATISTICS", "SYSTEM GLOBAL AREA"),
    "<PGA_chart>": ("PGA",),
    "<MEMORY_chart>": ("MEMORY STATISTICS",),
    "<log_switch_chart>": (
        "TẦN SUẤT LOG SWITCH",
        ("ONLINE REDO LOG", ("CẤU HÌNH LƯU TRỮ",)),
    ),
    "<aas_per_wait_class_for_instance_1>": (
        "WAIT CLASS",
        "AAS PER WAIT CLASS FOR INSTANCE 1",
    ),
    "<aas_per_wait_class_for_instance_1_chart>": (
        "WAIT CLASS",
        "AAS PER WAIT CLASS FOR INSTANCE 1",
    ),
    "<ash_top_timed_events_for_instance_1_for_9_days_of_history>": (
        "WAIT EVENT",
        "WAIT EVEN",
        "ASH TOP TIMED EVENTS FOR INSTANCE 1 FOR 9 DAYS OF HISTORY",
    ),
    "<ash_top_timed_events_for_instance_1_for_9_days_of_history_chart>": (
        "WAIT EVENT",
        "WAIT EVEN",
        "ASH TOP TIMED EVENTS FOR INSTANCE 1 FOR 9 DAYS OF HISTORY",
    ),
    "{{table_ash_top_sql_cluster_9_days}}": (
        "CÁC CÂU LỆNH TỐN NHIỀU DB TIME",
        "CAC CAU LENH TON NHIEU DB TIME",
        "ASH TOP SQL FOR CLUSTER",
    ),
    "{{table_ash_cpu_per_source}}": (
        "CÁC CÂU LỆNH SỬ DỤNG NHIỀU CPU",
        "CAC CAU LENH SU DUNG NHIEU CPU",
        "ASH CPU PER SOURCE",
    ),
    "<ash_cpu_per_source_chart>": (
        "CÁC CÂU LỆNH SỬ DỤNG NHIỀU CPU",
        "CAC CAU LENH SU DUNG NHIEU CPU",
        "ASH CPU PER SOURCE",
    ),
    "<ash_top_sql_for_instance_1_for_9_days_of_history_chart>": (
        "CÁC CÂU LỆNH TỐN NHIỀU DB TIME",
        "CAC CAU LENH TON NHIEU DB TIME",
        "ASH TOP SQL FOR INSTANCE 1 FOR 9 DAYS OF HISTORY",
        "ASH TOP SQL FOR CLUSTER",
    ),    "<log_switch_frequency_for_instance_2_chart>": (
        "Instances 2: Log switch frequency",
        "Instance 2: Log switch frequency",
        "LOG SWITCH FREQUENCY FOR INSTANCE 2",
    ),
    "<aas_per_wait_class_for_instance_2_chart>": (
        "Instance 2: AAS per Wait Class",
        "AAS PER WAIT CLASS FOR INSTANCE 2",
    ),
    "<ash_top_timed_events_for_instance_2_for_days_of_history_chart>": (
        "Instance 2: ASH Top Timed Events",
        "ASH TOP TIMED EVENTS FOR INSTANCE 2",
    ),
    "<ash_top_sql_for_instance_2_for_days_of_history_chart>": (
        "ASH TOP SQL FOR INSTANCE 2",
        "Instance 2: ASH Top SQL",
        "{{table_ash_top_sql_cluster_9_days}}",
    ),
    "<cpu_busy_and_idle_times_percent_for_instance_2_chart>": (
        "Instance 2: CPU Busy and Idle Times Percent",
        "CPU BUSY AND IDLE TIMES PERCENT FOR INSTANCE 2",
    ),
    "<memory_statistics_for_instance_2_chart>": (
        "Memory Statistics trên Instance 2",
        "MEMORY STATISTICS FOR INSTANCE 2",
    ),
    "<sga_statistics_for_instance_2_chart>": (
        "SGA Statistics trên Instance 2",
        "SGA STATISTICS FOR INSTANCE 2",
    ),
    "<pga_statistics_for_instance_2_chart>": (
        "PGA Statistics trên Instance 2",
        "PGA STATISTICS FOR INSTANCE 2",
    ),
}
