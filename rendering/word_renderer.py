from __future__ import annotations

from collections.abc import Callable
from pathlib import Path
import re
from time import perf_counter

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from models import ExtractedContent, GenerationReport, ImageContent, MappingRule, MultiImageContent, OperationCancelled, TableContent


LogCallback = Callable[[str], None]
CancelCheck = Callable[[], bool]


def render_report(
    template_path: str | Path,
    output_path: str | Path,
    resolved: dict[str, tuple[MappingRule, ExtractedContent]],
    rules: list[MappingRule],
    report: GenerationReport,
    log_callback: LogCallback | None = None,
    text_mapping: dict[str, str] | None = None,
    max_table_rows: int | None = None,
    lightweight_tables: bool = False,
    cleanup_unresolved_placeholders: bool = False,
    slow_step_seconds: float = 10.0,
    cancel_check: CancelCheck | None = None,
) -> None:
    log = _make_logger(log_callback)
    render_started = perf_counter()
    log(f"Starting Word report render: {output_path}")
    doc = Document(template_path)
    log("Word template loaded.")
    _raise_if_cancelled(cancel_check)

    # Thay thế {{collection_date}} nếu có trong template (bao gồm cả body, header, footer)
    # Mục đích: nhiều template đặt placeholder trong footer/header nên cần xử lý
    # cả các section.header/footer chứ không chỉ body.
    
    from datetime import datetime
    collection_date = (text_mapping or {}).get("{{collection_date}}") or datetime.now().strftime("%b-%Y")
    replaced_collection_date = False
    # Thay thế trong body
    for paragraph in iter_paragraphs(doc):
        _raise_if_cancelled(cancel_check)
        if "{{collection_date}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{collection_date}}", collection_date)
            replaced_collection_date = True
    # Thay thế trong header và footer
    for section in doc.sections:
        # Header
        # Header: duyệt các paragraph trong header của từng section
        for paragraph in section.header.paragraphs:
            _raise_if_cancelled(cancel_check)
            if "{{collection_date}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{collection_date}}", collection_date)
                replaced_collection_date = True
        # Footer
        # Footer: duyệt các paragraph trong footer của từng section
        for paragraph in section.footer.paragraphs:
            _raise_if_cancelled(cancel_check)
            if "{{collection_date}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{collection_date}}", collection_date)
                replaced_collection_date = True

    known_placeholders = set(resolved) | {rule.placeholder for rule in rules}
    present_placeholders = collect_present_placeholders(doc, known_placeholders)
    skipped_resolved = [placeholder for placeholder in resolved if placeholder not in present_placeholders]
    if skipped_resolved:
        log(f"Skipping resolved mappings not present in template: {len(skipped_resolved)}")

    for placeholder, (rule, content) in resolved.items():
        _raise_if_cancelled(cancel_check)
        if placeholder not in present_placeholders:
            report.missing_placeholders.append(f"{placeholder}: placeholder not found in template")
            continue
        step_started = perf_counter()
        content_for_render = _prepare_content_for_render(placeholder, content, max_table_rows, log)
        log(_describe_render_step("Rendering", placeholder, content_for_render))
        inserted = replace_placeholder(
            doc,
            placeholder,
            rule,
            content_for_render,
            lightweight_tables=lightweight_tables,
        )
        elapsed = perf_counter() - step_started
        if inserted:
            report.inserted.append(f"{placeholder} <- {content_for_render.source_path.name}")
            log(f"Completed placeholder {placeholder} in {elapsed:.2f}s")
        else:
            report.missing_placeholders.append(f"{placeholder}: placeholder not found in template")
            log(f"Placeholder not found {placeholder} after {elapsed:.2f}s")
        if elapsed >= slow_step_seconds:
            log(f"Slow render step: {placeholder} took {elapsed:.2f}s")

    for rule in rules:
        _raise_if_cancelled(cancel_check)
        if rule.placeholder not in resolved:
            if rule.placeholder not in present_placeholders:
                report.missing_placeholders.append(f"{rule.placeholder}: configured but not found in template")

    if cleanup_unresolved_placeholders:
        _raise_if_cancelled(cancel_check)
        cleanup_count = remove_unresolved_placeholder_paragraphs(doc, angle_only=True)
        if cleanup_count:
            log(f"Removed unresolved placeholder paragraphs: {cleanup_count}")

    if text_mapping:
        _raise_if_cancelled(cancel_check)
        replaced_count = replace_text_placeholders(doc, text_mapping)
        log(f"Replaced text placeholders: {replaced_count}")

    if cleanup_unresolved_placeholders:
        _raise_if_cancelled(cancel_check)
        cleanup_count = remove_unresolved_placeholder_paragraphs(doc, angle_only=False)
        if cleanup_count:
            log(f"Removed unresolved text placeholder paragraphs: {cleanup_count}")

    save_started = perf_counter()
    _raise_if_cancelled(cancel_check)
    log(f"Starting DOCX save: {output_path}")
    doc.save(output_path)
    save_elapsed = perf_counter() - save_started
    log(f"Finished DOCX save in {save_elapsed:.2f}s")
    if save_elapsed >= slow_step_seconds:
        log(f"Slow DOCX save: {save_elapsed:.2f}s")
    log(f"Finished Word report render in {perf_counter() - render_started:.2f}s")


def _raise_if_cancelled(cancel_check: CancelCheck | None) -> None:
    if cancel_check and cancel_check():
        raise OperationCancelled("Operation cancelled.")


def document_contains_placeholder(doc: DocumentObject, placeholder: str) -> bool:
    # Kiểm tra xem placeholder có tồn tại trong body hay trong các cell của table
    return any(placeholder in paragraph.text for paragraph in iter_document_paragraphs(doc))


def collect_present_placeholders(doc: DocumentObject, placeholders: set[str]) -> set[str]:
    if not placeholders:
        return set()
    document_text = "\n".join(paragraph.text for paragraph in iter_document_paragraphs(doc))
    return {placeholder for placeholder in placeholders if placeholder in document_text}


def replace_text_placeholders(doc: DocumentObject, mapping: dict[str, str]) -> int:
    replaced = 0
    for paragraph in iter_document_paragraphs(doc):
        for placeholder, replacement in mapping.items():
            if placeholder not in paragraph.text:
                continue
            _replace_text_in_paragraph(paragraph, placeholder, replacement)
            replaced += 1
    return replaced


def remove_unresolved_placeholder_paragraphs(doc: DocumentObject, angle_only: bool = False) -> int:
    removed = 0
    for paragraph in list(iter_document_paragraphs(doc)):
        if _contains_unresolved_placeholder(paragraph.text, angle_only=angle_only):
            previous = _previous_paragraph(paragraph)
            if previous is not None and _is_optional_content_heading(previous.text):
                remove_paragraph(previous)
                removed += 1
            remove_paragraph(paragraph)
            removed += 1
    return removed


def _contains_unresolved_placeholder(text: str, angle_only: bool = False) -> bool:
    if not text:
        return False
    if angle_only:
        return bool(re.fullmatch(r"\s*<[^<>\n]+>\s*", text))
    return bool(re.fullmatch(r"\s*\{\{[^{}]+\}\}\s*", text))


def _previous_paragraph(paragraph: Paragraph) -> Paragraph | None:
    previous = paragraph._element.getprevious()
    while previous is not None and previous.tag != qn("w:p"):
        previous = previous.getprevious()
    if previous is None:
        return None
    return Paragraph(previous, paragraph._parent)


def _is_optional_content_heading(text: str) -> bool:
    stripped = " ".join((text or "").split())
    if not stripped or ":" in stripped or len(stripped) > 90:
        return False
    letters = [character for character in stripped if character.isalpha()]
    if not letters:
        return False
    return stripped == stripped.upper()




def replace_placeholder(
    doc: DocumentObject,
    placeholder: str,
    rule: MappingRule,
    content: ExtractedContent,
    lightweight_tables: bool = False,
) -> bool:
    # Duyệt tất cả paragraph (cả trong tables) để tìm placeholder
    for paragraph in iter_paragraphs(doc):
        if placeholder not in paragraph.text:
            continue

        # Xoá nội dung paragraph hiện tại trước khi chèn nội dung mới
        clear_paragraph(paragraph)
        if isinstance(content, TableContent):
            insert_table_after_paragraph(
                doc,
                paragraph,
                content,
                lightweight=lightweight_tables,
                header_vertical=rule.table_header_vertical,
            )
        elif isinstance(content, MultiImageContent):
            insert_multi_image_after_paragraph(doc, paragraph, content, rule.width_inches)
        elif isinstance(content, ImageContent):
            insert_image_after_paragraph(doc, paragraph, content, rule.width_inches)
        else:
            paragraph.text = ""
        remove_paragraph(paragraph)
        return True
    return False


def iter_paragraphs(parent: DocumentObject | _Cell):
    # Trả về từng Paragraph trong một DocumentObject hoặc _Cell.
    # Lưu ý: nếu một cell chứa table lồng nhau thì hàm này cũng sẽ đệ quy.
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def iter_document_paragraphs(doc: DocumentObject):
    yield from iter_paragraphs(doc)
    for section in doc.sections:
        yield from iter_paragraphs(section.header)
        yield from iter_paragraphs(section.footer)


def _replace_text_in_paragraph(paragraph: Paragraph, placeholder: str, replacement: str) -> None:
    replaced_in_run = False
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)
            format_text_run(run)
            replaced_in_run = True

    if not replaced_in_run and placeholder in paragraph.text:
        paragraph.text = paragraph.text.replace(placeholder, replacement)
        for run in paragraph.runs:
            format_text_run(run)


def format_text_run(run) -> None:
    run.font.name = "Cambria"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")


def clear_paragraph(paragraph: Paragraph) -> None:
    # Xoá nội dung text trong mọi run của paragraph (giữ nguyên style/run)
    for run in paragraph.runs:
        run.text = ""


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    # Lưu ý: remove_paragraph dùng thao tác trực tiếp lên XML tree để gỡ phần tử paragraph


def insert_table_after_paragraph(
    doc: DocumentObject,
    paragraph: Paragraph,
    content: TableContent,
    lightweight: bool = False,
    header_vertical: bool = False,
) -> None:
    # Tạo table mới và chèn ngay sau paragraph chứa placeholder
    rows = content.rows
    if not rows and content.no_rows_selected:
        rows = [["No rows selected"]]
    if not rows:
        rows = [["No data found"]]

    row_count = len(rows)
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=row_count, cols=col_count)
    table.style = "Table Grid"

    for row_index, row_data in enumerate(rows):
        for col_index in range(col_count):
            cell = table.cell(row_index, col_index)
            cell.text = row_data[col_index] if col_index < len(row_data) else ""
            if lightweight:
                if row_index == 0 and row_count > 1:
                    format_cell(cell, is_header=True, vertical_text=header_vertical)
                continue
            # Định dạng cell: font, size, màu; và tô màu header nếu là hàng đầu
            is_header = row_index == 0 and row_count > 1
            format_cell(cell, is_header=is_header, vertical_text=is_header and header_vertical)

    set_table_autofit(table)
    paragraph._p.addnext(table._tbl)


def insert_image_after_paragraph(
    doc: DocumentObject,
    paragraph: Paragraph,
    content: ImageContent,
    width_inches: float | None,
) -> None:
    # Chèn ảnh vào một paragraph mới rồi đặt paragraph đó ngay sau paragraph gốc
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = image_paragraph.add_run()
    kwargs = {}
    if width_inches:
        kwargs["width"] = Inches(width_inches)
    run.add_picture(str(content.image_path), **kwargs)
    paragraph._p.addnext(image_paragraph._p)


def insert_multi_image_after_paragraph(
    doc: DocumentObject,
    paragraph: Paragraph,
    content: MultiImageContent,
    width_inches: float | None,
) -> None:
    insertions = []
    for index, image in enumerate(content.images):
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = image_paragraph.add_run()
        kwargs = {}
        if width_inches:
            kwargs["width"] = Inches(width_inches)
        run.add_picture(str(image.image_path), **kwargs)
        insertions.append(image_paragraph._p)

        caption = content.captions[index] if index < len(content.captions) else ""
        if caption:
            caption_paragraph = doc.add_paragraph()
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_paragraph.add_run(caption)
            caption_run.italic = True
            caption_run.font.name = "Cambria"
            caption_run.font.size = Pt(12)
            caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
            insertions.append(caption_paragraph._p)

    anchor = paragraph._p
    for element in reversed(insertions):
        anchor.addnext(element)


def format_cell(cell: _Cell, is_header: bool = False, vertical_text: bool = False) -> None:
    # Đặt font và kích thước chuẩn cho nội dung trong cell
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Cambria"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")

    if is_header:
        # Nếu là header: tô nền và đổi chữ sang màu trắng, chữ in đậm
        shading = parse_xml(r'<w:shd {} w:fill="0066CC"/>'.format(nsdecls("w")))
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            if vertical_text:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        if vertical_text:
            set_cell_text_direction(cell, "tbRl")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_autofit(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "autofit")
    tbl_pr.append(tbl_layout)


def set_cell_text_direction(cell: _Cell, direction: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:textDirection")):
        tc_pr.remove(existing)
    text_direction = OxmlElement("w:textDirection")
    text_direction.set(qn("w:val"), direction)
    tc_pr.append(text_direction)


def _make_logger(log_callback: LogCallback | None) -> LogCallback:
    def log(message: str) -> None:
        print(message)
        if log_callback:
            log_callback(message)

    return log


def _prepare_content_for_render(
    placeholder: str,
    content: ExtractedContent,
    max_table_rows: int | None,
    log: LogCallback,
) -> ExtractedContent:
    if not isinstance(content, TableContent) or max_table_rows is None:
        return content

    row_count = len(content.rows)
    if row_count <= max_table_rows:
        return content

    keep_rows = max(1, max_table_rows)
    truncated_rows = content.rows[:keep_rows]
    source_rows = max(0, row_count - 1)
    kept_data_rows = max(0, keep_rows - 1)
    log(
        f"Large table truncated for {placeholder}: "
        f"{source_rows} data rows -> {kept_data_rows} data rows"
    )
    return TableContent(
        source_path=content.source_path,
        rows=truncated_rows,
        title=content.title,
        section=content.section,
        logical_key=content.logical_key,
        keys=set(content.keys),
        index=content.index,
        no_rows_selected=content.no_rows_selected,
    )


def _describe_render_step(prefix: str, placeholder: str, content: ExtractedContent) -> str:
    source_name = content.title or content.logical_key or content.source_path.name
    if isinstance(content, TableContent):
        row_count = len(content.rows)
        col_count = max((len(row) for row in content.rows), default=0)
        data_rows = max(0, row_count - 1)
        return (
            f"{prefix} placeholder {placeholder}: "
            f"sheet/source '{source_name}', table, {data_rows} data rows, {col_count} columns"
        )
    if isinstance(content, ImageContent):
        return f"{prefix} placeholder {placeholder}: source '{source_name}', {content.content_type}"
    return f"{prefix} placeholder {placeholder}: source '{source_name}', {content.content_type}"
    # Ghi chú: thiết lập autofit giúp table co giãn theo nội dung/khổ trang
