from __future__ import annotations

import shutil
import hashlib
import sqlite3
import traceback
import uuid
import zipfile
from datetime import datetime, timedelta
from pathlib import Path, PurePosixPath, PureWindowsPath

from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
from openpyxl import Workbook

from ai_review import generate_ai_review
from app_logic import (
    DEFAULT_EDB360_MASTER_TEMPLATE,
    DEFAULT_SQL_MASTER_TEMPLATE,
    generate_edb360_report_to_file,
    generate_report_to_file,
    insert_placeholders_into_word,
    run_sql_one_click_pipeline,
    run_sql_pipeline,
)
from config import DEFAULT_MAPPING, load_mapping_rules
from app_logic import DEFAULT_SQL_MAPPING
from placeholder_manager import (
    add_placeholder_items,
    build_placeholder_from_token,
    delete_placeholder_item,
    extract_docx_placeholders,
    infer_placeholder_type,
    load_placeholder_items,
    placeholder_to_key,
    scan_placeholder_file,
    upsert_placeholder_item,
)


PROJECT_ROOT = Path(__file__).resolve().parents[1]
WEB_ROOT = Path(__file__).resolve().parent
RUNTIME_JOBS_DIR = PROJECT_ROOT / "runtime_jobs"
DATA_DIR = PROJECT_ROOT / "data"
UPLOADS_DIR = DATA_DIR / "uploads"
OUTPUTS_DIR = DATA_DIR / "outputs"
LOGS_DIR = DATA_DIR / "logs"
DB_PATH = DATA_DIR / "report_history.db"
MAX_JOB_AGE = timedelta(hours=6)

MODE_LABELS = {
    "oraclehc": "OracleHC",
    "edb360": "EDB360 One-click",
    "sqlhealthcheck": "SQLHealthcheck One-click",
}
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
ZIP_MEDIA_TYPE = "application/zip"
SUPPORTED_PLACEHOLDER_SCAN_SUFFIXES = {".docx", ".pdf", ".html", ".htm"}
SUPPORTED_AI_REVIEW_SUFFIXES = {".zip", ".docx", ".html", ".htm"}
REPORT_TYPES = {
    "01. Bao cao dinh ki",
    "02. Bao cao Health Check",
    "03. Bao cao Tuning",
    "04. Bao cao bao mat",
    "Unclassified",
}
PERIOD_TYPES = {"Quarter", "Month", "Custom", ""}
QUARTERS = {"Q1", "Q2", "Q3", "Q4", ""}
REPORT_STATUSES = {"Draft", "Generated", "Reviewed", "Sent", ""}

app = FastAPI(title="OracleHC Report Generator - Web")
app.mount("/static", StaticFiles(directory=WEB_ROOT / "static"), name="static")
app.mount("/assets", StaticFiles(directory=PROJECT_ROOT / "assets"), name="assets")
templates = Jinja2Templates(directory=WEB_ROOT / "templates")


@app.on_event("startup")
def on_startup() -> None:
    RUNTIME_JOBS_DIR.mkdir(parents=True, exist_ok=True)
    for folder in (DATA_DIR, UPLOADS_DIR, OUTPUTS_DIR, LOGS_DIR):
        folder.mkdir(parents=True, exist_ok=True)
    init_db()
    cleanup_old_jobs()


@app.get("/", response_class=HTMLResponse)
async def index(request: Request) -> HTMLResponse:
    return render_index(request)


@app.post("/generate")
async def generate_report_endpoint(
    request: Request,
    mode: str = Form(...),
    output_name: str = Form("final_healthcheck_report.docx"),
    template_file: UploadFile = File(...),
    source_zip: UploadFile = File(...),
):
    cleanup_old_jobs()
    job_dir = create_job_dir()
    input_dir = job_dir / "input"
    source_dir = job_dir / "source"
    output_dir = job_dir / "output"
    logs_dir = job_dir / "logs"
    for folder in (input_dir, source_dir, output_dir, logs_dir):
        folder.mkdir(parents=True, exist_ok=True)

    runtime_log = logs_dir / "runtime.log"
    error_log = logs_dir / "error.log"

    def log(message: str) -> None:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        with runtime_log.open("a", encoding="utf-8") as handle:
            handle.write(f"[{timestamp}] {message}\n")

    try:
        log("Job started")
        safe_mode = validate_mode(mode)
        safe_output_name = normalize_output_name(output_name)
        template_path = input_dir / "template.docx"
        zip_path = input_dir / "source.zip"
        output_path = output_dir / safe_output_name

        validate_upload_extension(template_file.filename, ".docx", "Word template")
        validate_upload_extension(source_zip.filename, ".zip", "Source package")

        log(f"Mode selected: {safe_mode}")
        await save_upload(template_file, template_path)
        log("Template saved")
        await save_upload(source_zip, zip_path)
        log("Source ZIP saved")

        safe_extract_zip(zip_path, source_dir)
        log("Source extracted")

        log("Generation started")
        generated_path = generate_web_report(
            mode=safe_mode,
            source_root=source_dir,
            template_path=template_path,
            output_path=output_path,
            log_callback=log,
        )
        log(f"Generation completed: {generated_path}")

        return FileResponse(
            generated_path,
            media_type=DOCX_MEDIA_TYPE,
            filename=safe_output_name,
        )
    except Exception as exc:
        error_log.write_text(traceback.format_exc(), encoding="utf-8")
        log(f"Generation failed: {exc}")
        return render_index(
            request,
            status_code=400,
            error_message="Generation failed.",
            error_reason=str(exc),
            selected_mode=mode if mode in MODE_LABELS else "oraclehc",
            output_name=output_name or "final_healthcheck_report.docx",
        )


@app.post("/api/template/scan")
async def scan_template_endpoint(
    mode: str = Form("oraclehc"),
    template_file: UploadFile = File(...),
) -> JSONResponse:
    safe_mode = validate_mode(mode)
    validate_upload_extension(template_file.filename, ".docx", "Word template")
    scan_dir = UPLOADS_DIR / "template_scans"
    scan_dir.mkdir(parents=True, exist_ok=True)
    scan_path = scan_dir / f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}.docx"
    await save_upload(template_file, scan_path)
    template_hash = hash_file(scan_path)
    result = scan_template_placeholders(scan_path, safe_mode)
    now = utc_now()
    try:
        persist_scan_rows(None, template_hash, result, now)
    except sqlite3.OperationalError as exc:
        if "disk I/O error" in str(exc):
            try:
                reset_db_after_io_error()
                persist_scan_rows(None, template_hash, result, now)
            except (sqlite3.OperationalError, OSError):
                result["scan_persisted"] = False
            else:
                result["scan_persisted"] = True
        else:
            result["scan_persisted"] = False
    else:
        result["scan_persisted"] = True
    result["template_name"] = template_file.filename
    result["template_hash"] = template_hash
    return JSONResponse(result)


@app.post("/api/template/insert")
async def insert_template_endpoint(
    template_file: UploadFile = File(...),
) -> FileResponse:
    validate_upload_extension(template_file.filename, ".docx", "Word template")
    insert_dir = UPLOADS_DIR / "template_inserts"
    insert_dir.mkdir(parents=True, exist_ok=True)
    template_path = insert_dir / f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}_{Path(template_file.filename).name}"
    await save_upload(template_file, template_path)
    insert_placeholders_into_word(template_path)
    download_name = f"{template_path.stem}_inserted.docx"
    return FileResponse(template_path, media_type=DOCX_MEDIA_TYPE, filename=download_name)


@app.post("/api/template/insert-job")
async def api_insert_template_job(
    background_tasks: BackgroundTasks,
    template_file: UploadFile = File(...),
) -> JSONResponse:
    validate_upload_extension(template_file.filename, ".docx", "Word template")

    job_id = create_job_id()
    upload_dir = UPLOADS_DIR / job_id
    output_dir = OUTPUTS_DIR / job_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)

    original_name = Path(template_file.filename).name
    template_path = upload_dir / "template.docx"
    output_name = f"{Path(original_name).stem}_inserted.docx"
    output_path = output_dir / output_name
    await save_upload(template_file, template_path)

    now = utc_now()
    with db_connect() as conn:
        conn.execute(
            """
            INSERT INTO report_jobs (
                id, mode, output_file_name, template_file_name, source_package_name,
                template_hash, source_hash, status, progress, current_step,
                created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                job_id,
                "oraclehc",
                output_name,
                original_name,
                "",
                hash_file(template_path),
                "",
                "processing",
                0,
                "Insert job created",
                now,
                now,
            ),
        )
    add_job_log(job_id, "info", "insert_created", "Placeholder insert job created")
    background_tasks.add_task(process_template_insert_job, job_id, template_path, output_path)
    return JSONResponse({"success": True, "job_id": job_id, "status": "processing"})


@app.get("/api/placeholders")
async def api_placeholders(mode: str = "oraclehc") -> JSONResponse:
    safe_mode = validate_mode(mode)
    mapping_path = mapping_path_for_mode(safe_mode)
    rows = load_placeholder_items(mapping_path)
    return JSONResponse(
        {
            "success": True,
            "mode": safe_mode,
            "mapping_file": str(mapping_path),
            "placeholders": rows,
            "summary": {
                "total": len(rows),
                "tables": sum(1 for item in rows if item.get("content_type") == "table"),
                "charts": sum(1 for item in rows if item.get("content_type") == "chart"),
                "other": sum(1 for item in rows if item.get("content_type") not in {"table", "chart"}),
            },
        }
    )


@app.post("/api/placeholders/scan")
async def api_placeholders_scan(
    mode: str = Form("oraclehc"),
    template_file: UploadFile = File(...),
) -> JSONResponse:
    safe_mode = validate_mode(mode)
    suffix = validate_placeholder_scan_extension(template_file.filename)
    scan_dir = UPLOADS_DIR / "placeholder_scans"
    scan_dir.mkdir(parents=True, exist_ok=True)
    scan_path = scan_dir / f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}{suffix}"
    await save_upload(template_file, scan_path)
    result = scan_placeholder_file(scan_path, mapping_path_for_mode(safe_mode))
    result["template_name"] = template_file.filename
    result["template_hash"] = hash_file(scan_path)
    return JSONResponse(result)


@app.post("/api/placeholders/add")
async def api_placeholders_add(request: Request) -> JSONResponse:
    payload = await request.json()
    safe_mode = validate_mode(str(payload.get("mode", "oraclehc")))
    raw_items = payload.get("placeholders") or []
    if not isinstance(raw_items, list):
        raise HTTPException(status_code=400, detail="placeholders must be a list")

    items = []
    for raw in raw_items:
        if isinstance(raw, str):
            items.append(build_placeholder_from_token(raw))
        elif isinstance(raw, dict):
            token = str(raw.get("placeholder", "")).strip()
            if not token:
                continue
            item = build_placeholder_from_token(token, str(raw.get("source_file") or raw.get("template") or ""))
            item.update({key: value for key, value in raw.items() if value not in (None, "")})
            items.append(item)

    added = add_placeholder_items(mapping_path_for_mode(safe_mode), items)
    return JSONResponse({"success": True, "added": added, "added_count": len(added)})


@app.put("/api/placeholders")
async def api_placeholders_upsert(request: Request) -> JSONResponse:
    payload = await request.json()
    safe_mode = validate_mode(str(payload.get("mode", "oraclehc")))
    item = payload.get("placeholder") or payload.get("item") or {}
    if not isinstance(item, dict):
        raise HTTPException(status_code=400, detail="placeholder must be an object")
    try:
        saved = upsert_placeholder_item(
            mapping_path_for_mode(safe_mode),
            item,
            str(payload.get("original_placeholder") or ""),
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return JSONResponse({"success": True, "placeholder": saved})


@app.delete("/api/placeholders")
async def api_placeholders_delete(request: Request) -> JSONResponse:
    payload = await request.json()
    safe_mode = validate_mode(str(payload.get("mode", "oraclehc")))
    placeholder = str(payload.get("placeholder", "")).strip()
    if not placeholder:
        raise HTTPException(status_code=400, detail="placeholder is required")
    deleted = delete_placeholder_item(mapping_path_for_mode(safe_mode), placeholder)
    if not deleted:
        raise HTTPException(status_code=404, detail="placeholder not found")
    return JSONResponse({"success": True, "deleted": placeholder})


def persist_scan_rows(job_id: str | None, template_hash: str, result: dict, created_at: str) -> None:
    with db_connect() as conn:
        for item in result["detected_placeholders"]:
            conn.execute(
                """
                INSERT INTO template_scan_results (
                    job_id, template_hash, placeholder, mapping_key, placeholder_type,
                    status, source, created_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    job_id,
                    template_hash,
                    item["placeholder"],
                    item["mapping_key"],
                    item["type"],
                    item["status"],
                    item["source"],
                    created_at,
                ),
            )


@app.post("/api/report/generate")
async def api_generate_report(
    background_tasks: BackgroundTasks,
    mode: str = Form(...),
    output_name: str = Form("final_healthcheck_report.docx"),
    template_file: UploadFile = File(...),
    source_zip: UploadFile = File(...),
) -> JSONResponse:
    safe_mode = validate_mode(mode)
    validate_upload_extension(template_file.filename, ".docx", "Word template")
    validate_upload_extension(source_zip.filename, ".zip", "Source package")
    safe_output_name = normalize_output_name(output_name)

    job_id = create_job_id()
    upload_dir = UPLOADS_DIR / job_id
    output_dir = OUTPUTS_DIR / job_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)

    template_path = upload_dir / "template.docx"
    zip_path = upload_dir / "source.zip"
    await save_upload(template_file, template_path)
    await save_upload(source_zip, zip_path)

    template_hash = hash_file(template_path)
    source_hash = hash_file(zip_path)
    now = utc_now()
    with db_connect() as conn:
        conn.execute(
            """
            INSERT INTO report_jobs (
                id, mode, output_file_name, template_file_name, source_package_name,
                template_hash, source_hash, status, progress, current_step,
                created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                job_id,
                safe_mode,
                safe_output_name,
                template_file.filename,
                source_zip.filename,
                template_hash,
                source_hash,
                "processing",
                0,
                "Job created",
                now,
                now,
            ),
        )
    add_job_log(job_id, "info", "job_created", "Job created")
    update_job(job_id, progress=10, current_step="Template uploaded")
    add_job_log(job_id, "success", "upload_template", "Template uploaded")
    update_job(job_id, progress=20, current_step="Source package uploaded")
    add_job_log(job_id, "success", "upload_source", "Source package uploaded")

    background_tasks.add_task(
        process_report_job,
        job_id,
        safe_mode,
        template_path,
        zip_path,
        output_dir / safe_output_name,
    )
    return JSONResponse({"success": True, "job_id": job_id, "status": "processing"})


@app.post("/api/report/generate-edb360")
async def api_generate_edb360_report(
    background_tasks: BackgroundTasks,
    output_name: str = Form("final_edb360_report.docx"),
    customer_name: str = Form(""),
    system_name: str = Form(""),
    database_name: str = Form(""),
    database_display_name: str = Form(""),
    creator: str = Form(""),
    approver: str = Form(""),
    version: str = Form(""),
    collection_date: str = Form(""),
    source_zip: UploadFile = File(...),
) -> JSONResponse:
    if not DEFAULT_EDB360_MASTER_TEMPLATE.is_file():
        raise HTTPException(status_code=500, detail="Internal EDB360 master template is missing")
    validate_upload_extension(source_zip.filename, ".zip", "EDB360 source package")
    safe_output_name = normalize_output_name(output_name)

    job_id = create_job_id()
    upload_dir = UPLOADS_DIR / job_id
    output_dir = OUTPUTS_DIR / job_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)

    zip_path = upload_dir / "source.zip"
    await save_upload(source_zip, zip_path)
    metadata = {
        "customer_name": customer_name,
        "system_name": system_name,
        "database_name": database_name,
        "database_display_name": database_display_name,
        "creator": creator,
        "approver": approver,
        "version": version,
        "collection_date": collection_date,
    }

    now = utc_now()
    with db_connect() as conn:
        conn.execute(
            """
            INSERT INTO report_jobs (
                id, mode, output_file_name, template_file_name, source_package_name,
                template_hash, source_hash, status, progress, current_step,
                created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                job_id,
                "edb360",
                safe_output_name,
                DEFAULT_EDB360_MASTER_TEMPLATE.name,
                source_zip.filename,
                hash_file(DEFAULT_EDB360_MASTER_TEMPLATE),
                hash_file(zip_path),
                "processing",
                0,
                "EDB360 job created",
                now,
                now,
            ),
        )
    add_job_log(job_id, "info", "job_created", "EDB360 one-click job created")
    update_job(job_id, progress=20, current_step="EDB360 source package uploaded")
    add_job_log(job_id, "success", "upload_source", "EDB360 source package uploaded")

    background_tasks.add_task(
        process_edb360_report_job,
        job_id,
        zip_path,
        output_dir / safe_output_name,
        metadata,
    )
    return JSONResponse({"success": True, "job_id": job_id, "status": "processing"})


@app.post("/api/report/generate-sqlhealthcheck")
async def api_generate_sqlhealthcheck_report(
    background_tasks: BackgroundTasks,
    output_name: str = Form("sqlhealthcheck_reports.zip"),
    creator: str = Form(""),
    approver: str = Form(""),
    version: str = Form(""),
    collection_date: str = Form(""),
    source_zip: UploadFile = File(...),
) -> JSONResponse:
    if not DEFAULT_SQL_MASTER_TEMPLATE.is_file():
        raise HTTPException(status_code=500, detail="Internal SQLHealthcheck master template is missing")
    validate_upload_extension(source_zip.filename, ".zip", "SQLHealthcheck source package")
    safe_output_name = normalize_zip_output_name(output_name, "sqlhealthcheck_reports.zip")

    job_id = create_job_id()
    upload_dir = UPLOADS_DIR / job_id
    output_dir = OUTPUTS_DIR / job_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)

    zip_path = upload_dir / "source.zip"
    await save_upload(source_zip, zip_path)
    metadata = {
        "creator": creator,
        "approver": approver,
        "version": version,
        "collection_date": collection_date,
    }

    now = utc_now()
    with db_connect() as conn:
        conn.execute(
            """
            INSERT INTO report_jobs (
                id, mode, output_file_name, template_file_name, source_package_name,
                template_hash, source_hash, status, progress, current_step,
                created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                job_id,
                "sqlhealthcheck",
                safe_output_name,
                DEFAULT_SQL_MASTER_TEMPLATE.name,
                source_zip.filename,
                hash_file(DEFAULT_SQL_MASTER_TEMPLATE),
                hash_file(zip_path),
                "processing",
                0,
                "SQLHealthcheck one-click job created",
                now,
                now,
            ),
        )
    add_job_log(job_id, "info", "job_created", "SQLHealthcheck one-click job created")
    update_job(job_id, progress=20, current_step="SQLHealthcheck source package uploaded")
    add_job_log(job_id, "success", "upload_source", "SQLHealthcheck source package uploaded")

    background_tasks.add_task(
        process_sqlhealthcheck_report_job,
        job_id,
        zip_path,
        output_dir / safe_output_name,
        metadata,
    )
    return JSONResponse({"success": True, "job_id": job_id, "status": "processing"})


@app.get("/api/report/status/{job_id}")
async def api_report_status(job_id: str) -> JSONResponse:
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    logs = get_job_logs(job_id)
    return JSONResponse({"success": True, "job": job, "logs": logs})


@app.get("/api/report/download/{job_id}")
async def api_report_download(job_id: str) -> FileResponse:
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job["status"] != "success" or not job["output_file_path"]:
        raise HTTPException(status_code=409, detail="Report is not ready")
    output_path = Path(job["output_file_path"])
    if not output_path.is_file():
        raise HTTPException(status_code=404, detail="Output file is missing")
    return FileResponse(output_path, media_type=media_type_for_path(output_path), filename=job["output_file_name"])


@app.post("/api/ai-review/generate")
async def api_ai_review_generate(
    review_type: str = Form("oracle_health_check"),
    output_style: str = Form("short_table"),
    history_job_id: str = Form(""),
    source_file: UploadFile | None = File(None),
) -> JSONResponse:
    if review_type != "oracle_health_check":
        raise HTTPException(status_code=400, detail="Only Oracle Health Check review is supported")
    if output_style != "short_table":
        raise HTTPException(status_code=400, detail="Only Short Table output is supported")

    review_id = create_job_id()
    review_dir = UPLOADS_DIR / "ai_reviews" / review_id
    review_dir.mkdir(parents=True, exist_ok=True)

    try:
        if history_job_id.strip():
            job = get_job(history_job_id.strip())
            if not job:
                raise HTTPException(status_code=404, detail="History source package not found")
            source_zip = UPLOADS_DIR / history_job_id.strip() / "source.zip"
            if not source_zip.is_file():
                raise HTTPException(status_code=404, detail="History source ZIP is missing")
            source_root = review_dir / "source"
            source_root.mkdir(parents=True, exist_ok=True)
            safe_extract_zip(source_zip, source_root)
            source_name = job.get("source_package_name") or "History source package"
        else:
            if source_file is None:
                raise HTTPException(status_code=400, detail="Select a History package or upload a source file")
            suffix = validate_ai_review_extension(source_file.filename)
            source_path = review_dir / f"source{suffix}"
            await save_upload(source_file, source_path)
            source_name = source_file.filename or source_path.name
            if suffix == ".zip":
                source_root = review_dir / "source"
                source_root.mkdir(parents=True, exist_ok=True)
                safe_extract_zip(source_path, source_root)
            else:
                source_root = source_path

        result = generate_ai_review(source_root)
        return JSONResponse(
            {
                "success": True,
                "review_id": review_id,
                "source_name": source_name,
                "used_ai": result["used_ai"],
                "provider": result["provider"],
                "rows": result["rows"],
                "debug": result["debug"],
            }
        )
    except HTTPException:
        raise
    except Exception as exc:
        debug_log = LOGS_DIR / f"ai_review_{review_id}.log"
        debug_log.parent.mkdir(parents=True, exist_ok=True)
        debug_log.write_text(traceback.format_exc(), encoding="utf-8")
        raise HTTPException(status_code=500, detail=f"AI Review failed: {exc}") from exc


@app.post("/api/ai-review/export/docx")
async def api_ai_review_export_docx(request: Request) -> FileResponse:
    rows = normalize_ai_review_rows(await request.json())
    export_dir = OUTPUTS_DIR / "ai_reviews"
    export_dir.mkdir(parents=True, exist_ok=True)
    output_path = export_dir / f"ai_review_{create_job_id()}.docx"
    document = Document()
    document.add_heading("AI Review", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Mục"
    headers[1].text = "Đánh giá"
    headers[2].text = "Khuyến nghị"
    for item in rows:
        cells = table.add_row().cells
        cells[0].text = item["section"]
        cells[1].text = item["assessment"]
        cells[2].text = item["recommendation"]
    document.save(output_path)
    return FileResponse(output_path, media_type=DOCX_MEDIA_TYPE, filename="ai_review.docx")


@app.post("/api/ai-review/export/xlsx")
async def api_ai_review_export_xlsx(request: Request) -> FileResponse:
    rows = normalize_ai_review_rows(await request.json())
    export_dir = OUTPUTS_DIR / "ai_reviews"
    export_dir.mkdir(parents=True, exist_ok=True)
    output_path = export_dir / f"ai_review_{create_job_id()}.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "AI Review"
    sheet.append(["Mục", "Đánh giá", "Khuyến nghị"])
    for item in rows:
        sheet.append([item["section"], item["assessment"], item["recommendation"]])
    for column in ("A", "B", "C"):
        sheet.column_dimensions[column].width = 34 if column == "A" else 58
    workbook.save(output_path)
    return FileResponse(output_path, media_type=XLSX_MEDIA_TYPE, filename="ai_review.xlsx")


@app.get("/api/history")
async def api_history() -> JSONResponse:
    with db_connect() as conn:
        rows = conn.execute(
            """
            SELECT report_jobs.*, companies.company_name, companies.short_name AS company_short_name
            FROM report_jobs
            LEFT JOIN companies ON companies.id = report_jobs.company_id
            ORDER BY report_jobs.created_at DESC
            LIMIT 200
            """
        ).fetchall()
    return JSONResponse({"success": True, "jobs": [row_to_dict(row) for row in rows]})


@app.get("/api/history/{job_id}")
async def api_history_detail(job_id: str) -> JSONResponse:
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return JSONResponse({"success": True, "job": job})


@app.put("/api/history/{job_id}/assignment")
async def api_history_assignment(job_id: str, request: Request) -> JSONResponse:
    if not get_job(job_id):
        raise HTTPException(status_code=404, detail="Job not found")
    payload = await request.json()
    assignment = normalize_assignment_payload(payload)
    company_id = assignment.pop("company_id")
    if company_id:
        company = get_company(company_id)
        if not company:
            raise HTTPException(status_code=400, detail="Company not found")

    with db_connect() as conn:
        conn.execute(
            """
            UPDATE report_jobs
            SET company_id = ?, report_type = ?, period_type = ?, quarter = ?, month = ?,
                year = ?, assessment_from = ?, assessment_to = ?, report_status = ?,
                note = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                company_id,
                assignment["report_type"],
                assignment["period_type"],
                assignment["quarter"],
                assignment["month"],
                assignment["year"],
                assignment["assessment_from"],
                assignment["assessment_to"],
                assignment["report_status"],
                assignment["note"],
                utc_now(),
                job_id,
            ),
        )
    return JSONResponse({"success": True, "job": get_job(job_id)})


@app.get("/api/companies")
async def api_companies() -> JSONResponse:
    with db_connect() as conn:
        rows = conn.execute(
            """
            SELECT
                companies.*,
                COUNT(report_jobs.id) AS total_reports,
                (
                    SELECT output_file_name
                    FROM report_jobs
                    WHERE report_jobs.company_id = companies.id
                    ORDER BY created_at DESC
                    LIMIT 1
                ) AS latest_report,
                COALESCE(MAX(report_jobs.updated_at), companies.updated_at) AS last_updated
            FROM companies
            LEFT JOIN report_jobs ON report_jobs.company_id = companies.id
            GROUP BY companies.id
            ORDER BY companies.updated_at DESC, companies.company_name COLLATE NOCASE
            """
        ).fetchall()
    return JSONResponse({"success": True, "companies": [row_to_dict(row) for row in rows]})


@app.post("/api/companies")
async def api_company_create(request: Request) -> JSONResponse:
    payload = await request.json()
    company = normalize_company_payload(payload)
    company_id = create_job_id()
    now = utc_now()
    with db_connect() as conn:
        conn.execute(
            """
            INSERT INTO companies (
                id, company_name, short_name, customer_code, contact_person,
                email, phone, note, created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                company_id,
                company["company_name"],
                company["short_name"],
                company["customer_code"],
                company["contact_person"],
                company["email"],
                company["phone"],
                company["note"],
                now,
                now,
            ),
        )
    return JSONResponse({"success": True, "company": get_company(company_id)})


@app.put("/api/companies/{company_id}")
async def api_company_update(company_id: str, request: Request) -> JSONResponse:
    if not get_company(company_id):
        raise HTTPException(status_code=404, detail="Company not found")
    payload = await request.json()
    company = normalize_company_payload(payload)
    with db_connect() as conn:
        conn.execute(
            """
            UPDATE companies
            SET company_name = ?, short_name = ?, customer_code = ?, contact_person = ?,
                email = ?, phone = ?, note = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                company["company_name"],
                company["short_name"],
                company["customer_code"],
                company["contact_person"],
                company["email"],
                company["phone"],
                company["note"],
                utc_now(),
                company_id,
            ),
        )
    return JSONResponse({"success": True, "company": get_company(company_id)})


@app.delete("/api/companies/{company_id}")
async def api_company_delete(company_id: str) -> JSONResponse:
    if not get_company(company_id):
        raise HTTPException(status_code=404, detail="Company not found")
    with db_connect() as conn:
        conn.execute("UPDATE report_jobs SET company_id = NULL, updated_at = ? WHERE company_id = ?", (utc_now(), company_id))
        conn.execute("DELETE FROM companies WHERE id = ?", (company_id,))
    return JSONResponse({"success": True})


@app.get("/api/companies/{company_id}/reports")
async def api_company_reports(
    company_id: str,
    report_type: str = Query("all"),
    quarter: str = Query("all"),
    year: str = Query("all"),
    report_status: str = Query("all"),
) -> JSONResponse:
    company = get_company(company_id)
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")
    clauses = ["report_jobs.company_id = ?"]
    values: list[str] = [company_id]
    filters = {
        "report_type": report_type,
        "quarter": quarter,
        "year": year,
        "report_status": report_status,
    }
    for field, value in filters.items():
        if value and value != "all":
            clauses.append(f"{field} = ?")
            values.append(value)
    with db_connect() as conn:
        rows = conn.execute(
            f"""
            SELECT report_jobs.*, companies.company_name, companies.short_name AS company_short_name
            FROM report_jobs
            LEFT JOIN companies ON companies.id = report_jobs.company_id
            WHERE {" AND ".join(clauses)}
            ORDER BY report_jobs.created_at DESC
            """,
            values,
        ).fetchall()
    return JSONResponse({"success": True, "company": company, "reports": [row_to_dict(row) for row in rows]})


@app.get("/api/history/{job_id}/logs")
async def api_history_logs(job_id: str) -> JSONResponse:
    if not get_job(job_id):
        raise HTTPException(status_code=404, detail="Job not found")
    return JSONResponse({"success": True, "logs": get_job_logs(job_id)})


@app.get("/api/history/{job_id}/placeholders")
async def api_history_placeholders(job_id: str) -> JSONResponse:
    with db_connect() as conn:
        rows = conn.execute(
            "SELECT * FROM template_scan_results WHERE job_id = ? ORDER BY id",
            (job_id,),
        ).fetchall()
    return JSONResponse({"success": True, "placeholders": [row_to_dict(row) for row in rows]})


@app.post("/api/history/{job_id}/regenerate")
async def api_history_regenerate(job_id: str) -> JSONResponse:
    raise HTTPException(status_code=501, detail="Regenerate will be connected after upload reuse is finalized.")


@app.delete("/api/history/{job_id}")
async def api_history_delete(job_id: str) -> JSONResponse:
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    with db_connect() as conn:
        conn.execute("DELETE FROM report_job_logs WHERE job_id = ?", (job_id,))
        conn.execute("DELETE FROM template_scan_results WHERE job_id = ?", (job_id,))
        conn.execute("DELETE FROM report_jobs WHERE id = ?", (job_id,))
    return JSONResponse({"success": True})


def render_index(
    request: Request,
    status_code: int = 200,
    error_message: str | None = None,
    error_reason: str | None = None,
    selected_mode: str = "oraclehc",
    output_name: str = "final_oraclehc_report.docx",
) -> HTMLResponse:
    return templates.TemplateResponse(
        request,
        "dashboard_v2.html",
        {
            "error_message": error_message,
            "error_reason": error_reason,
            "selected_mode": selected_mode,
            "output_name": output_name,
        },
        status_code=status_code,
    )


def create_job_dir() -> Path:
    job_id = create_job_id()
    job_dir = RUNTIME_JOBS_DIR / job_id
    job_dir.mkdir(parents=True, exist_ok=False)
    return job_dir


def create_job_id() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:8]


def cleanup_old_jobs(max_age: timedelta = MAX_JOB_AGE) -> None:
    RUNTIME_JOBS_DIR.mkdir(parents=True, exist_ok=True)
    cutoff = datetime.now() - max_age
    for job_dir in RUNTIME_JOBS_DIR.iterdir():
        if not job_dir.is_dir():
            continue
        try:
            modified_at = datetime.fromtimestamp(job_dir.stat().st_mtime)
            if modified_at < cutoff:
                shutil.rmtree(job_dir)
        except OSError:
            continue


def validate_mode(mode: str) -> str:
    normalized = mode.strip().lower()
    if normalized not in MODE_LABELS:
        raise ValueError(f"Unsupported mode: {mode}")
    return normalized


def validate_upload_extension(filename: str | None, expected_suffix: str, label: str) -> None:
    if not filename:
        raise ValueError(f"{label} is required.")
    if Path(filename).suffix.lower() != expected_suffix:
        raise ValueError(f"{label} must be a {expected_suffix} file.")


def validate_placeholder_scan_extension(filename: str | None) -> str:
    if not filename:
        raise ValueError("Placeholder scan file is required.")
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_PLACEHOLDER_SCAN_SUFFIXES:
        raise ValueError("Placeholder scan file must be .docx, .pdf, .html, or .htm.")
    return suffix


def validate_ai_review_extension(filename: str | None) -> str:
    if not filename:
        raise HTTPException(status_code=400, detail="AI Review source file is required")
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_AI_REVIEW_SUFFIXES:
        raise HTTPException(status_code=400, detail="AI Review source must be .zip, .docx, .html, or .htm")
    return suffix


def normalize_ai_review_rows(payload: dict) -> list[dict[str, str]]:
    raw_rows = payload.get("rows") if isinstance(payload, dict) else None
    if not isinstance(raw_rows, list):
        raise HTTPException(status_code=400, detail="rows must be a list")
    rows = []
    for item in raw_rows:
        if not isinstance(item, dict):
            continue
        rows.append(
            {
                "section": str(item.get("section", "")).strip(),
                "assessment": str(item.get("assessment", "")).strip(),
                "recommendation": str(item.get("recommendation", "")).strip(),
            }
        )
    if not rows:
        raise HTTPException(status_code=400, detail="No AI Review rows to export")
    return rows


def mapping_path_for_mode(mode: str) -> Path:
    return DEFAULT_SQL_MAPPING if mode == "sqlhealthcheck" else DEFAULT_MAPPING


def normalize_output_name(output_name: str) -> str:
    clean_name = Path(output_name.strip() or "final_healthcheck_report.docx").name
    if not clean_name:
        clean_name = "final_healthcheck_report.docx"
    if Path(clean_name).suffix.lower() != ".docx":
        clean_name = f"{Path(clean_name).stem or 'final_healthcheck_report'}.docx"
    return clean_name


def normalize_zip_output_name(output_name: str, default_name: str) -> str:
    clean_name = Path(output_name.strip() or default_name).name
    if not clean_name:
        clean_name = default_name
    if Path(clean_name).suffix.lower() != ".zip":
        clean_name = f"{Path(clean_name).stem or Path(default_name).stem}.zip"
    return clean_name


def media_type_for_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        return XLSX_MEDIA_TYPE
    if suffix == ".zip":
        return ZIP_MEDIA_TYPE
    return DOCX_MEDIA_TYPE


async def save_upload(upload: UploadFile, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with destination.open("wb") as handle:
        while chunk := await upload.read(1024 * 1024):
            handle.write(chunk)


def safe_extract_zip(zip_path: Path, dest_dir: Path) -> None:
    dest_root = dest_dir.resolve()
    with zipfile.ZipFile(zip_path, "r") as archive:
        for member in archive.infolist():
            filename = member.filename
            posix_path = PurePosixPath(filename)
            windows_path = PureWindowsPath(filename)
            if posix_path.is_absolute() or windows_path.is_absolute():
                raise ValueError(f"Unsafe ZIP path detected: {filename}")
            if ".." in posix_path.parts or ".." in windows_path.parts:
                raise ValueError(f"Unsafe ZIP path detected: {filename}")
            resolved_path = (dest_dir / filename).resolve()
            if not resolved_path.is_relative_to(dest_root):
                raise ValueError(f"Unsafe ZIP path detected: {filename}")
        archive.extractall(dest_dir)


def generate_web_report(
    mode: str,
    source_root: Path,
    template_path: Path,
    output_path: Path,
    log_callback,
) -> str:
    if mode == "oraclehc":
        return generate_report_to_file(
            html_input=source_root,
            word_file=template_path,
            output_file=output_path,
            mapping_file=DEFAULT_MAPPING,
            chart_output_dir=output_path.parent / "generated_charts",
            log_callback=log_callback,
        )

    if mode == "sqlhealthcheck":
        generated_files = run_sql_pipeline(
            input_root=source_root,
            template_file=template_path,
            output_root=output_path.parent,
            log_callback=log_callback,
        )
        generated_docx = next(
            (Path(path) for path in generated_files if Path(path).suffix.lower() == ".docx"),
            None,
        )
        if generated_docx is None or not generated_docx.is_file():
            raise ValueError("SQLHealthcheck did not produce a Word report.")
        if generated_docx.resolve() != output_path.resolve():
            shutil.copyfile(generated_docx, output_path)
        return str(output_path)

    raise ValueError(f"Unsupported mode: {mode}")


def process_report_job(
    job_id: str,
    mode: str,
    template_path: Path,
    zip_path: Path,
    output_path: Path,
) -> None:
    started_at = datetime.now()
    source_dir = UPLOADS_DIR / job_id / "source"
    log_file = LOGS_DIR / f"{job_id}.log"

    try:
        update_job(job_id, progress=30, current_step="Extracting source package")
        add_job_log(job_id, "info", "extract_source", "Extracting source package")
        source_dir.mkdir(parents=True, exist_ok=True)
        safe_extract_zip(zip_path, source_dir)
        add_job_log(job_id, "success", "extract_source", "Source package extracted")

        update_job(job_id, progress=45, current_step="Scanning placeholders")
        scan_result = scan_template_placeholders(template_path, mode)
        save_scan_result(job_id, hash_file(template_path), scan_result)
        add_job_log(
            job_id,
            "success",
            "scan_template",
            f"Detected {scan_result['summary']['total']} placeholders from Word template",
        )

        update_job(job_id, progress=60, current_step=f"Parsing {mode_label(mode)} data")
        add_job_log(job_id, "info", "parse_source", f"Parsing {mode_label(mode)} source data")

        update_job(job_id, progress=75, current_step="Replacing placeholders")
        add_job_log(job_id, "info", "render_report", "Replacing placeholders and rendering report")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        generated_path = generate_web_report(
            mode=mode,
            source_root=source_dir,
            template_path=template_path,
            output_path=output_path,
            log_callback=lambda message: add_job_log(job_id, "info", "generator", message),
        )

        update_job(job_id, progress=90, current_step="Building final Word report")
        final_path = Path(generated_path)
        if final_path.resolve() != output_path.resolve():
            shutil.copyfile(final_path, output_path)
        add_job_log(job_id, "success", "export_report", "Final Word report built")

        duration = (datetime.now() - started_at).total_seconds()
        update_job(
            job_id,
            status="success",
            progress=100,
            current_step="Report ready",
            output_file_path=str(output_path),
            log_file_path=str(log_file),
            duration_seconds=duration,
            completed_at=utc_now(),
        )
        add_job_log(job_id, "success", "complete", "Report ready")
    except Exception as exc:
        error_text = str(exc)
        log_file.write_text(traceback.format_exc(), encoding="utf-8")
        update_job(
            job_id,
            status="failed",
            current_step="Generation failed",
            error_message=error_text,
            log_file_path=str(log_file),
            completed_at=utc_now(),
        )
        add_job_log(job_id, "error", "failed", error_text)


def process_edb360_report_job(
    job_id: str,
    zip_path: Path,
    output_path: Path,
    metadata: dict[str, str],
) -> None:
    started_at = datetime.now()
    source_dir = UPLOADS_DIR / job_id / "source"
    log_file = LOGS_DIR / f"{job_id}.log"

    try:
        update_job(job_id, progress=35, current_step="Extracting EDB360 package")
        add_job_log(job_id, "info", "extract_source", "Extracting EDB360 package")
        source_dir.mkdir(parents=True, exist_ok=True)
        safe_extract_zip(zip_path, source_dir)
        add_job_log(job_id, "success", "extract_source", "EDB360 package extracted")

        update_job(job_id, progress=50, current_step="Scanning internal master template")
        scan_result = scan_template_placeholders(DEFAULT_EDB360_MASTER_TEMPLATE, "oraclehc")
        save_scan_result(job_id, hash_file(DEFAULT_EDB360_MASTER_TEMPLATE), scan_result)
        add_job_log(
            job_id,
            "success",
            "scan_template",
            f"Detected {scan_result['summary']['total']} placeholders from internal master template",
        )

        update_job(job_id, progress=65, current_step="Parsing EDB360 data")
        add_job_log(job_id, "info", "parse_source", "Parsing EDB360 HTML data")

        update_job(job_id, progress=78, current_step="Rendering Word from master template")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        generated_path = generate_edb360_report_to_file(
            html_input=source_dir,
            output_file=output_path,
            metadata=metadata,
            chart_output_dir=output_path.parent / "generated_charts",
            log_callback=lambda message: add_job_log(job_id, "info", "generator", message),
        )

        update_job(job_id, progress=92, current_step="Finalizing Word report")
        final_path = Path(generated_path)
        if final_path.resolve() != output_path.resolve():
            shutil.copyfile(final_path, output_path)
        add_job_log(job_id, "success", "export_report", "Final EDB360 Word report built")

        duration = (datetime.now() - started_at).total_seconds()
        update_job(
            job_id,
            status="success",
            progress=100,
            current_step="Report ready",
            output_file_path=str(output_path),
            log_file_path=str(log_file),
            duration_seconds=duration,
            completed_at=utc_now(),
        )
        add_job_log(job_id, "success", "complete", "Report ready")
    except Exception as exc:
        error_text = str(exc)
        log_file.write_text(traceback.format_exc(), encoding="utf-8")
        update_job(
            job_id,
            status="failed",
            current_step="EDB360 generation failed",
            error_message=error_text,
            log_file_path=str(log_file),
            completed_at=utc_now(),
        )
        add_job_log(job_id, "error", "failed", error_text)


def process_sqlhealthcheck_report_job(
    job_id: str,
    zip_path: Path,
    output_path: Path,
    metadata: dict[str, str],
) -> None:
    started_at = datetime.now()
    source_dir = UPLOADS_DIR / job_id / "source"
    work_dir = OUTPUTS_DIR / job_id / "sqlhealthcheck_work"
    log_file = LOGS_DIR / f"{job_id}.log"

    try:
        update_job(job_id, progress=35, current_step="Extracting SQLHealthcheck package")
        add_job_log(job_id, "info", "extract_source", "Extracting SQLHealthcheck package")
        source_dir.mkdir(parents=True, exist_ok=True)
        safe_extract_zip(zip_path, source_dir)
        add_job_log(job_id, "success", "extract_source", "SQLHealthcheck package extracted")

        update_job(job_id, progress=50, current_step="Scanning internal SQLHealthcheck template")
        scan_result = scan_template_placeholders(DEFAULT_SQL_MASTER_TEMPLATE, "sqlhealthcheck")
        save_scan_result(job_id, hash_file(DEFAULT_SQL_MASTER_TEMPLATE), scan_result)
        add_job_log(
            job_id,
            "success",
            "scan_template",
            f"Detected {scan_result['summary']['total']} placeholders from internal SQLHealthcheck template",
        )

        update_job(job_id, progress=65, current_step="Merging SQLHealthcheck CSV files")
        add_job_log(job_id, "info", "parse_source", "Merging SQLHealthcheck CSV data")
        work_dir.mkdir(parents=True, exist_ok=True)
        generated_files = run_sql_one_click_pipeline(
            input_root=source_dir,
            output_root=work_dir,
            metadata=metadata,
            log_callback=lambda message: add_job_log(job_id, "info", "generator", message),
        )

        update_job(job_id, progress=90, current_step="Packaging SQLHealthcheck output")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        package_generated_files(generated_files, work_dir, output_path)
        add_job_log(job_id, "success", "export_report", "SQLHealthcheck Word and Excel files packaged")

        duration = (datetime.now() - started_at).total_seconds()
        update_job(
            job_id,
            status="success",
            progress=100,
            current_step="Report package ready",
            output_file_path=str(output_path),
            log_file_path=str(log_file),
            duration_seconds=duration,
            completed_at=utc_now(),
        )
        add_job_log(job_id, "success", "complete", "Report package ready")
    except Exception as exc:
        error_text = str(exc)
        log_file.write_text(traceback.format_exc(), encoding="utf-8")
        update_job(
            job_id,
            status="failed",
            current_step="SQLHealthcheck generation failed",
            error_message=error_text,
            log_file_path=str(log_file),
            completed_at=utc_now(),
        )
        add_job_log(job_id, "error", "failed", error_text)


def package_generated_files(generated_files: list[str], base_dir: Path, output_path: Path) -> None:
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for file_name in generated_files:
            path = Path(file_name)
            if not path.is_file():
                continue
            try:
                archive_name = path.relative_to(base_dir)
            except ValueError:
                archive_name = Path(path.name)
            archive.write(path, archive_name.as_posix())


def process_template_insert_job(job_id: str, template_path: Path, output_path: Path) -> None:
    started_at = datetime.now()
    log_file = LOGS_DIR / f"{job_id}.log"

    try:
        update_job(job_id, progress=20, current_step="Preparing Word template")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(template_path, output_path)
        add_job_log(job_id, "success", "prepare_template", "Template copied for placeholder insertion")

        update_job(job_id, progress=45, current_step="Loading placeholder mapping")
        report = insert_placeholders_into_word(
            output_path,
            log_callback=lambda message: add_job_log(job_id, "info", "placeholder_insert", message),
        )

        update_job(job_id, progress=90, current_step="Finalizing inserted template")
        add_job_log(
            job_id,
            "success",
            "placeholder_insert",
            f"Inserted {len(report.inserted)} placeholders; {len(report.missing_anchors)} anchors missing",
        )

        duration = (datetime.now() - started_at).total_seconds()
        update_job(
            job_id,
            status="success",
            progress=100,
            current_step="Inserted template ready",
            output_file_path=str(output_path),
            log_file_path=str(log_file),
            duration_seconds=duration,
            completed_at=utc_now(),
        )
        add_job_log(job_id, "success", "complete", "Inserted template ready")
    except Exception as exc:
        error_text = str(exc)
        log_file.write_text(traceback.format_exc(), encoding="utf-8")
        update_job(
            job_id,
            status="failed",
            current_step="Placeholder insertion failed",
            error_message=error_text,
            log_file_path=str(log_file),
            completed_at=utc_now(),
        )
        add_job_log(job_id, "error", "failed", error_text)

def init_db() -> None:
    schema = """
            CREATE TABLE IF NOT EXISTS companies (
                id TEXT PRIMARY KEY,
                company_name TEXT NOT NULL,
                short_name TEXT NOT NULL,
                customer_code TEXT,
                contact_person TEXT,
                email TEXT,
                phone TEXT,
                note TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS report_jobs (
                id TEXT PRIMARY KEY,
                mode TEXT NOT NULL,
                output_file_name TEXT NOT NULL,
                template_file_name TEXT,
                source_package_name TEXT,
                template_hash TEXT,
                source_hash TEXT,
                status TEXT NOT NULL,
                progress INTEGER DEFAULT 0,
                current_step TEXT,
                output_file_path TEXT,
                log_file_path TEXT,
                error_message TEXT,
                duration_seconds REAL,
                company_id TEXT,
                report_type TEXT DEFAULT 'Unclassified',
                period_type TEXT,
                quarter TEXT,
                month TEXT,
                year TEXT,
                assessment_from TEXT,
                assessment_to TEXT,
                report_status TEXT DEFAULT 'Generated',
                note TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                completed_at TEXT,
                FOREIGN KEY(company_id) REFERENCES companies(id)
            );

            CREATE TABLE IF NOT EXISTS report_job_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                job_id TEXT NOT NULL,
                timestamp TEXT NOT NULL,
                level TEXT NOT NULL,
                step TEXT,
                message TEXT NOT NULL,
                FOREIGN KEY(job_id) REFERENCES report_jobs(id)
            );

            CREATE TABLE IF NOT EXISTS template_scan_results (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                job_id TEXT,
                template_hash TEXT,
                placeholder TEXT NOT NULL,
                mapping_key TEXT,
                placeholder_type TEXT,
                status TEXT,
                source TEXT,
                created_at TEXT NOT NULL
            );
            """
    try:
        with db_connect() as conn:
            conn.executescript(schema)
            migrate_db(conn)
    except sqlite3.OperationalError as exc:
        if "disk I/O error" not in str(exc):
            raise
        quarantine_db_files()
        with db_connect() as conn:
            conn.executescript(schema)
            migrate_db(conn)


def migrate_db(conn: sqlite3.Connection) -> None:
    existing = {row["name"] for row in conn.execute("PRAGMA table_info(report_jobs)").fetchall()}
    columns = {
        "company_id": "TEXT",
        "report_type": "TEXT DEFAULT 'Unclassified'",
        "period_type": "TEXT",
        "quarter": "TEXT",
        "month": "TEXT",
        "year": "TEXT",
        "assessment_from": "TEXT",
        "assessment_to": "TEXT",
        "report_status": "TEXT DEFAULT 'Generated'",
        "note": "TEXT",
    }
    for name, definition in columns.items():
        if name not in existing:
            conn.execute(f"ALTER TABLE report_jobs ADD COLUMN {name} {definition}")
    conn.execute("UPDATE report_jobs SET report_type = 'Unclassified' WHERE report_type IS NULL OR report_type = ''")
    conn.execute("UPDATE report_jobs SET report_status = 'Generated' WHERE report_status IS NULL OR report_status = ''")


def db_connect() -> sqlite3.Connection:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH, timeout=30)
    conn.execute("PRAGMA busy_timeout = 30000")
    conn.execute("PRAGMA journal_mode = MEMORY")
    conn.row_factory = sqlite3.Row
    return conn


def reset_db_after_io_error() -> None:
    quarantine_db_files()
    init_db()


def quarantine_db_files() -> None:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    for path in (DB_PATH, Path(f"{DB_PATH}-journal")):
        if path.exists():
            try:
                path.rename(path.with_name(f"{path.name}.corrupt_{stamp}"))
            except OSError:
                continue


def update_job(job_id: str, **fields) -> None:
    if not fields:
        return
    fields["updated_at"] = utc_now()
    assignments = ", ".join(f"{name} = ?" for name in fields)
    values = list(fields.values()) + [job_id]
    with db_connect() as conn:
        conn.execute(f"UPDATE report_jobs SET {assignments} WHERE id = ?", values)


def add_job_log(job_id: str, level: str, step: str, message: str) -> None:
    timestamp = utc_now()
    with db_connect() as conn:
        conn.execute(
            """
            INSERT INTO report_job_logs (job_id, timestamp, level, step, message)
            VALUES (?, ?, ?, ?, ?)
            """,
            (job_id, timestamp, level, step, message),
        )
    log_file = LOGS_DIR / f"{job_id}.log"
    log_file.parent.mkdir(parents=True, exist_ok=True)
    with log_file.open("a", encoding="utf-8") as handle:
        handle.write(f"[{timestamp}] [{level}] [{step}] {message}\n")


def get_job(job_id: str) -> dict | None:
    with db_connect() as conn:
        row = conn.execute(
            """
            SELECT report_jobs.*, companies.company_name, companies.short_name AS company_short_name
            FROM report_jobs
            LEFT JOIN companies ON companies.id = report_jobs.company_id
            WHERE report_jobs.id = ?
            """,
            (job_id,),
        ).fetchone()
    return row_to_dict(row) if row else None


def get_company(company_id: str) -> dict | None:
    with db_connect() as conn:
        row = conn.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    return row_to_dict(row) if row else None


def normalize_company_payload(payload: dict) -> dict:
    company_name = str(payload.get("company_name") or payload.get("companyName") or "").strip()
    short_name = str(payload.get("short_name") or payload.get("shortName") or "").strip()
    if not company_name:
        raise HTTPException(status_code=400, detail="Company Name is required")
    if not short_name:
        raise HTTPException(status_code=400, detail="Short Name is required")
    return {
        "company_name": company_name,
        "short_name": short_name,
        "customer_code": clean_optional(payload.get("customer_code") or payload.get("customerCode")),
        "contact_person": clean_optional(payload.get("contact_person") or payload.get("contactPerson")),
        "email": clean_optional(payload.get("email")),
        "phone": clean_optional(payload.get("phone")),
        "note": clean_optional(payload.get("note")),
    }


def normalize_assignment_payload(payload: dict) -> dict:
    report_type = clean_optional(payload.get("report_type") or payload.get("reportType")) or "Unclassified"
    period_type = clean_optional(payload.get("period_type") or payload.get("periodType"))
    quarter = clean_optional(payload.get("quarter"))
    report_status = clean_optional(payload.get("report_status") or payload.get("reportStatus")) or "Generated"
    if report_type not in REPORT_TYPES:
        raise HTTPException(status_code=400, detail="Unsupported report type")
    if period_type not in PERIOD_TYPES:
        raise HTTPException(status_code=400, detail="Unsupported period type")
    if quarter not in QUARTERS:
        raise HTTPException(status_code=400, detail="Unsupported quarter")
    if report_status not in REPORT_STATUSES:
        raise HTTPException(status_code=400, detail="Unsupported report status")
    return {
        "company_id": clean_optional(payload.get("company_id") or payload.get("companyId")),
        "report_type": report_type,
        "period_type": period_type,
        "quarter": quarter,
        "month": clean_optional(payload.get("month")),
        "year": clean_optional(payload.get("year")),
        "assessment_from": clean_optional(payload.get("assessment_from") or payload.get("assessmentFrom")),
        "assessment_to": clean_optional(payload.get("assessment_to") or payload.get("assessmentTo")),
        "report_status": report_status,
        "note": clean_optional(payload.get("note")),
    }


def clean_optional(value) -> str:
    return str(value or "").strip()


def get_job_logs(job_id: str) -> list[dict]:
    with db_connect() as conn:
        rows = conn.execute(
            "SELECT * FROM report_job_logs WHERE job_id = ? ORDER BY id",
            (job_id,),
        ).fetchall()
    return [row_to_dict(row) for row in rows]


def row_to_dict(row: sqlite3.Row) -> dict:
    return {key: row[key] for key in row.keys()}


def utc_now() -> str:
    return datetime.now().isoformat(timespec="seconds")


def hash_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def mode_label(mode: str) -> str:
    return MODE_LABELS.get(mode, mode)


def scan_template_placeholders(template_path: Path, mode: str) -> dict:
    placeholders = extract_docx_placeholders(template_path)
    registry = placeholder_registry(mode)
    seen: set[str] = set()
    duplicate_set = {placeholder for placeholder in placeholders if placeholders.count(placeholder) > 1}
    detected = []

    for placeholder in placeholders:
        if placeholder in seen:
            continue
        seen.add(placeholder)
        mapping_key = placeholder_to_key(placeholder)
        rule = registry.get(placeholder) or registry.get(f"{{{{{mapping_key}}}}}") or registry.get(mapping_key)
        status = "mapped" if rule else "missing_mapping"
        if placeholder in duplicate_set:
            status = "duplicate"
        detected.append(
            {
                "placeholder": placeholder,
                "mapping_key": mapping_key,
                "type": (rule or {}).get("type", infer_placeholder_type(mapping_key)),
                "status": status,
                "source": (rule or {}).get("source", source_name_for_mode(mode)),
            }
        )

    summary = {
        "total": len(detected),
        "mapped": sum(1 for item in detected if item["status"] == "mapped"),
        "missing_mapping": sum(1 for item in detected if item["status"] == "missing_mapping"),
        "duplicates": sum(1 for item in detected if item["status"] == "duplicate"),
    }
    return {"success": True, "detected_placeholders": detected, "summary": summary}


def source_name_for_mode(mode: str) -> str:
    return "sqlhealthcheck_csv" if mode == "sqlhealthcheck" else "oraclehc_html"


def placeholder_registry(mode: str) -> dict[str, dict]:
    mapping_path = DEFAULT_SQL_MAPPING if mode == "sqlhealthcheck" else DEFAULT_MAPPING
    registry: dict[str, dict] = {}
    try:
        for rule in load_mapping_rules(mapping_path):
            mapping_key = rule.source_key or placeholder_to_key(rule.placeholder)
            info = {
                "type": rule.content_type,
                "source": source_name_for_mode(mode),
            }
            registry[rule.placeholder] = info
            registry[mapping_key] = info
    except Exception:
        pass

    required_tables = {
        "{{table_ash_top_sql_cluster_9_days}}": {
            "type": "table",
            "source": source_name_for_mode(mode),
        },
        "{{table_ash_cpu_per_source}}": {
            "type": "table",
            "source": source_name_for_mode(mode),
        },
    }
    registry.update(required_tables)
    registry["table_ash_top_sql_cluster_9_days"] = required_tables["{{table_ash_top_sql_cluster_9_days}}"]
    registry["table_ash_cpu_per_source"] = required_tables["{{table_ash_cpu_per_source}}"]
    return registry


def save_scan_result(job_id: str, template_hash: str, scan_result: dict) -> None:
    now = utc_now()
    with db_connect() as conn:
        for item in scan_result["detected_placeholders"]:
            conn.execute(
                """
                INSERT INTO template_scan_results (
                    job_id, template_hash, placeholder, mapping_key, placeholder_type,
                    status, source, created_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    job_id,
                    template_hash,
                    item["placeholder"],
                    item["mapping_key"],
                    item["type"],
                    item["status"],
                    item["source"],
                    now,
                ),
            )
