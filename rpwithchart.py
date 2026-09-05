from __future__ import annotations

import os
import re
from collections.abc import Callable, Iterable
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph


LogCallback = Callable[[str], None]
DEFAULT_SQL_MAX_RENDER_ROWS = int(os.getenv("SQLHEALCHECK_MAX_RENDER_ROWS", "200"))


SQL_TABLE_MAPPING = {
    "<volume_info>": {
        "sheet": "Volume Info",
        "columns": [0, 1, 2, 3, 4, 5],
        "transpose": True,
        "key_value_table": True,
        "font_size": 12,
        "header_height": 0.65,
        "row_height": 0.65,
        "row_height_rule": "atLeast",
        "column_widths": [7.6, 8.8],
    },
    "<file_size>": {
        "sheet": "File Sizes and Space",
        "columns": [0, 1, 2, 3, 4, 5, 7],
        "max_rows": 50,
        "font_size": 12,
        "column_widths": [2.9, 6.5, 1.8, 1.9, 1.1, 2.0, 2.2],
    },
    "<fileio>": {
        "sheet": "IO Stats By File",
        "max_rows": 50,
        "vertical_header": True,
        "vertical_body": True,
        "horizontal_columns": ["Database Name", "Logical Name", "type_desc", "Physical Name", "file_id"],
        "header_height": 2.6,
        "dynamic_vertical_row_height": True,
        "min_row_height": 1.8,
        "max_row_height": 3.4,
        "row_height_rule": "atLeast",
        "font_size": 12,
        "column_widths": [
            2.5,
            2.8,
            0.9,
            1.0,
            7.7,
            0.8,
            0.85,
            0.85,
            0.85,
            0.85,
            0.85,
            0.85,
            0.85,
            0.85,
            0.85,
            0.85,
            0.85,
            0.85,
            0.85,
            0.85,
        ],
    },
    "<conn_count>": {
        "sheet": "Connection Counts by IP Address",
        "max_rows": 50,
    },
    "<cpu_usage>": {
        "sheet": "CPU Usage by Database",
        "columns": [0, 1, 3],
        "max_rows": 50,
        "font_size": 12,
        "column_widths": [5.0, 6.2, 5.2],
    },
    "<io_usage>": {
        "sheet": "IO Usage By Database",
        "columns": [0, 1, 3],
        "max_rows": 50,
        "font_size": 12,
        "column_widths": [5.0, 6.2, 5.2],
    },
    "<buffer_usage>": {
        "sheet": "Total Buffer Usage by Database",
        "columns": [0, 1, 3],
        "max_rows": 50,
        "font_size": 12,
        "column_widths": [5.0, 6.2, 5.2],
    },
    "<top_worker>": {
        "sheet": "Top Worker Time Queries",
        "columns": [0, 1, 2, 4],
        "max_rows": 50,
        "font_size": 12,
        "column_widths": [6.4, 7.7, 6.3, 6.3],
    },
    "<missing_index>": {
        "sheet": "Missing Indexes",
        "columns": [2, 5, 6, 7, 9],
        "max_rows": 50,
        "font_size": 12,
        "column_widths": [3.0, 4.0, 4.0, 4.0, 4.0],
    },
    "<agent_job>": {
        "sheet": "SQL Server Agent Jobs",
        "columns": [0, 1, 2, 3, 4, 8, 9],
        "max_rows": 50,
        "font_size": 12,
    },
    "<recent_bk>": {
        "sheet": "Recent Full Backups",
        "columns": [2, 3, 4, 5, 11],
        "max_rows": 50,
    },
    "<collect_date>": {},
}


SQL_CHART_MAPPING = {
    "<cpu_chart>": {
        "sheet": "CPU Usage by Database",
        "title": "Chart 1. CPU Usage by Database",
        "label_col": 1,
        "value_col": 3,
        "top_n": 10,
    },
    "<io_chart>": {
        "sheet": "IO Usage By Database",
        "title": "Chart 2. IO Usage By Database",
        "label_col": 1,
        "value_col": 3,
        "top_n": 10,
    },
    "<buffer_chart>": {
        "sheet": "Total Buffer Usage by Database",
        "title": "Chart 3. Total Buffer Usage by Database",
        "label_col": 1,
        "value_col": 3,
        "top_n": 10,
    },
}


SQL_SHEET_ALIASES = {
    "Missing Indexes": ["Missing Indexes", "Missing Indexes All Databases"],
    "Recent Full Backups": ["Recent Full Backups", "Last Backup By Database"],
    "Last Backup By Database": ["Last Backup By Database", "Recent Full Backups"],
}


SQL_SCALAR_MAPPING = {
    "<volume_inf_per>": {
        "sheet": "Volume Info",
        "value_column": "Space Free %",
        "selector_column": "Space Free %",
        "selector": "min",
        "format": "{:.2f}",
    },
    "<volume_inf_size>": {
        "sheet": "Volume Info",
        "value_column": "Available Size (GB)",
        "selector_column": "Space Free %",
        "selector": "min",
        "format": "{:.2f}",
    },
}


def render_excel_report(
    excel_path: str | Path,
    template_path: str | Path,
    output_path: str | Path,
    mapping_path: str | Path | None = None,
    text_mapping: dict[str, str] | None = None,
    log_callback: LogCallback | None = None,
    max_table_rows: int | None = DEFAULT_SQL_MAX_RENDER_ROWS,
    lightweight_tables: bool = True,
    slow_step_seconds: float = 10.0,
) -> str:
    """Render SQLHealcheck Excel data into the Word template using SQL-specific mappings."""
    del mapping_path, lightweight_tables, slow_step_seconds

    excel = Path(excel_path)
    template = Path(template_path)
    output = Path(output_path)
    log = _make_logger(log_callback)

    if not excel.is_file():
        raise FileNotFoundError(f"Excel input does not exist: {excel}")
    if not template.is_file():
        raise FileNotFoundError(f"Word template does not exist: {template}")
    if template.suffix.lower() != ".docx":
        raise ValueError(f"Word template must be a .docx file: {template}")
    if output.suffix.lower() != ".docx":
        raise ValueError(f"Output report must be a .docx file: {output}")

    output.parent.mkdir(parents=True, exist_ok=True)
    log(f"SQLHealcheck Excel input: {excel}")
    log(f"SQLHealcheck template: {template}")
    log("Rendering SQLHealcheck Word report with fixed SQL placeholder mapping...")

    generated = generate_sql_healthcheck_report(
        excel_file=excel,
        template_file=template,
        output_file=output,
        mapping=SQL_TABLE_MAPPING,
        chart_mapping=SQL_CHART_MAPPING,
        scalar_mapping=SQL_SCALAR_MAPPING,
        text_mapping=text_mapping,
        log_callback=log_callback,
        max_table_rows=max_table_rows,
    )
    log(f"Word report created: {generated}")
    return str(generated)


def generate_sql_healthcheck_report(
    excel_file: str | Path,
    template_file: str | Path,
    output_file: str | Path,
    mapping: dict,
    chart_mapping: dict | None = None,
    scalar_mapping: dict | None = None,
    text_mapping: dict[str, str] | None = None,
    log_callback: LogCallback | None = None,
    max_table_rows: int | None = DEFAULT_SQL_MAX_RENDER_ROWS,
) -> str:
    log = _make_logger(log_callback)
    excel_path = Path(excel_file)
    output_path = Path(output_file)
    xls = pd.ExcelFile(excel_path)
    doc = Document(template_file)
    temp_images: list[Path] = []
    report_context = build_sql_report_context(xls, text_mapping or {})
    apply_sql_cover_metadata(doc, report_context, log)

    for placeholder, config in mapping.items():
        try:
            if placeholder == "<collect_date>":
                current_date = datetime.now().strftime("%m.%Y")
                replaced = replace_placeholder_text(doc, placeholder, current_date)
                log(f"Replaced {placeholder} with {current_date}" if replaced else f"Placeholder not found: {placeholder}")
                continue

            if not config:
                continue

            sheet_name = resolve_sql_sheet_name(xls, config["sheet"])
            if not sheet_name:
                log(f"Could not process {placeholder}: sheet not found '{config['sheet']}'")
                continue

            dataframe = pd.read_excel(xls, sheet_name=sheet_name)
            dataframe = _prepare_dataframe(dataframe, config)

            max_rows = config.get("max_rows")
            if max_rows is None and max_table_rows is not None:
                max_rows = max_table_rows
            if max_rows and len(dataframe) > max_rows:
                dataframe = dataframe.head(max_rows)

            inserted = replace_placeholder_with_table(doc, placeholder, dataframe, config)
            if inserted:
                log(f"Replaced {placeholder} with sheet '{sheet_name}' (rows={len(dataframe)})")
            else:
                log(f"Placeholder not found: {placeholder}")

        except Exception as exc:
            log(f"Could not process {placeholder}: {exc}")

    if chart_mapping:
        for placeholder, config in chart_mapping.items():
            try:
                sheet_name = resolve_sql_sheet_name(xls, config["sheet"])
                if not sheet_name:
                    log(f"Could not create chart for {placeholder}: sheet not found '{config['sheet']}'")
                    continue

                dataframe = pd.read_excel(xls, sheet_name=sheet_name)
                temp_image = output_path.parent / f"temp_chart_{placeholder.strip('<>').replace('_', '')}.png"
                temp_images.append(temp_image)

                if create_pie_chart(
                    dataframe,
                    config.get("title", sheet_name),
                    temp_image,
                    config.get("label_col", 0),
                    config.get("value_col", 1),
                    config.get("top_n", 10),
                    log,
                ):
                    inserted = replace_placeholder_with_image(doc, placeholder, temp_image)
                    log(f"Inserted chart for {placeholder}" if inserted else f"Placeholder not found: {placeholder}")
            except Exception as exc:
                log(f"Could not create chart for {placeholder}: {exc}")

    if scalar_mapping:
        for placeholder, config in scalar_mapping.items():
            try:
                value = extract_scalar_value(xls, config)
                replaced = replace_placeholder_text(doc, placeholder, value)
                log(f"Replaced {placeholder} with {value}" if replaced else f"Placeholder not found: {placeholder}")
            except Exception as exc:
                log(f"Could not process {placeholder}: {exc}")

    try:
        update_sql_hardware_table(doc, xls, log)
    except Exception as exc:
        log(f"Could not update SQLHealthcheck hardware sections: {exc}")

    try:
        update_sql_summary_table(doc, xls, log, report_context.get("report_language", "vi"))
    except Exception as exc:
        log(f"Could not update SQLHealthcheck summary section: {exc}")

    try:
        doc.save(output_path)
        log(f"Report generated: {output_path}")
    except PermissionError:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = output_path.with_name(f"{output_path.stem}_{timestamp}{output_path.suffix}")
        doc.save(output_path)
        log(f"File is open. Saved as: {output_path}")
    finally:
        for temp_image in temp_images:
            try:
                temp_image.unlink(missing_ok=True)
            except Exception:
                pass
        try:
            xls.close()
        except Exception:
            pass

    return str(output_path)


def build_sql_report_context(xls: pd.ExcelFile, text_mapping: dict[str, str]) -> dict[str, str]:
    server_row = _first_sheet_row(xls, "Server Properties")
    version_row = _first_sheet_row(xls, "Version Info")
    server_name = _clean_text(server_row.get("ServerName") or version_row.get("Server Name") or "")
    machine_name = _clean_text(server_row.get("MachineName") or _machine_from_server_name(server_name))
    return {
        "server_name": text_mapping.get("server_name") or server_name,
        "machine_name": text_mapping.get("machine_name") or machine_name,
        "creator": text_mapping.get("creator") or "Trần Đinh Nhất Đăng",
        "approver": text_mapping.get("approver") or "Hồ Quốc Trí",
        "version": text_mapping.get("version") or "1.0",
        "collection_date": text_mapping.get("collection_date") or datetime.now().strftime("%m.%Y"),
        "report_language": _normalize_report_language(text_mapping.get("report_language")),
    }


def apply_sql_cover_metadata(doc: DocumentObject, context: dict[str, str], log: LogCallback) -> None:
    _update_sql_cover_paragraph(doc, context)

    replacements = {
        "Phan Ngoc Phuong Nhi": context["creator"],
        "Phan Ngọc Phương Nhi": context["creator"],
        "Nhi, Phan Ngoc Phuong": context["creator"],
        "Nhi, Phan Ngọc Phương": context["creator"],
        "Ho Quoc Tri": context["approver"],
        "Tri, Ho Quoc": context["approver"],
        "Hồ Quốc Trí": context["approver"],
        "{{creator}}": context["creator"],
        "{{approver}}": context["approver"],
        "{{version}}": context["version"],
        "{{collection_date}}": context["collection_date"],
        "{{server_name}}": context.get("server_name", ""),
        "{{cover_server_name}}": context.get("server_name", ""),
    }
    if context.get("server_name"):
        replacements.update(
            {
                "DC-SQL06\\DB KHTT": context["server_name"],
                "DC-SQL06\\DBKHTT": context["server_name"],
                "DC-SQL06\\DBCF": context["server_name"],
            }
        )

    for old_value, new_value in replacements.items():
        if old_value and new_value:
            replace_placeholder_text(doc, old_value, new_value)

    if context.get("version"):
        _replace_value_after_label(doc, "Version", context["version"])
    if context.get("collection_date"):
        replace_placeholder_text(doc, "<collect_date>", context["collection_date"])
    updated_cover_cells = update_sql_cover_metadata_tables(doc, context)
    if updated_cover_cells:
        log(f"Updated SQLHealthcheck cover metadata cells: {updated_cover_cells}")
    updated_approval_cells = update_sql_approval_tables(doc, context)
    if updated_approval_cells:
        log(f"Updated SQLHealthcheck approval metadata cells: {updated_approval_cells}")
    log("Applied SQLHealthcheck cover metadata")


def update_sql_cover_metadata_tables(doc: DocumentObject, context: dict[str, str]) -> int:
    values = {
        "author": context["creator"],
        "version": context["version"],
        "collection date": context["collection_date"],
    }
    updated = 0
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = _cell_text(row.cells[0]).lower()
            if label in values and values[label]:
                _set_cell_text_preserve_format(row.cells[1], values[label])
                updated += 1
    return updated


def update_sql_approval_tables(doc: DocumentObject, context: dict[str, str]) -> int:
    updated = 0
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        headers = [_cell_text(cell).lower() for cell in table.rows[0].cells]
        if "name" not in headers or "position" not in headers:
            continue
        name_index = headers.index("name")
        position_index = headers.index("position")
        for row in table.rows[1:]:
            if len(row.cells) <= max(name_index, position_index):
                continue
            position = _cell_text(row.cells[position_index]).lower()
            if "database engineer" in position:
                _set_cell_text_preserve_format(row.cells[name_index], context["creator"])
                updated += 1
            elif "database solution manager" in position or "solution manager" in position:
                _set_cell_text_preserve_format(row.cells[name_index], context["approver"])
                updated += 1
    return updated


def _update_sql_cover_paragraph(doc: DocumentObject, context: dict[str, str]) -> bool:
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text
        if "MAINTENANCE REPORT" not in text or "Author" not in text:
            continue
        updated = text
        if context.get("server_name"):
            updated = re.sub(
                r"(MAINTENANCE REPORT\s+).*?(?=\s+Author)",
                lambda match: f"{match.group(1)}{context['server_name']}",
                updated,
                flags=re.DOTALL,
            )
        updated = re.sub(r"(Author\s+).*?(?=\s+Version)", lambda match: f"{match.group(1)}{context['creator']}", updated, flags=re.DOTALL)
        updated = re.sub(r"(Version\s+).*?(?=\s+Collection Date)", lambda match: f"{match.group(1)}{context['version']}", updated, flags=re.DOTALL)
        updated = re.sub(
            r"(Collection Date\s+).*?(?=\s+MỤC LỤC|$)",
            lambda match: f"{match.group(1)}{context['collection_date']}",
            updated,
            flags=re.DOTALL,
        )
        _set_paragraph_text(paragraph, updated)
        return True
    return False


def update_sql_hardware_table(doc: DocumentObject, xls: pd.ExcelFile, log: LogCallback) -> None:
    values = {
        **build_sql_volume_values(xls),
        **build_sql_hardware_values(xls),
    }
    if not values:
        log("No SQLHealthcheck static table data found")
        return

    placeholder_updates = replace_sql_static_placeholders(doc, values)
    updated = 0
    styled = 0
    for table in doc.tables:
        labels = [_cell_text(row.cells[0]) for row in table.rows if row.cells]
        if _is_sql_volume_table(labels):
            updated += update_sql_volume_table(table, build_sql_volume_rows(xls))
            style_sql_key_value_table(table)
            styled += 1
            continue
        matching_labels = [label for label in labels if label in values]
        if len(matching_labels) < 2:
            continue
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = _cell_text(row.cells[0])
            if label in values and values[label]:
                _set_cell_text(row.cells[1], values[label])
                updated += 1
        style_sql_key_value_table(table)
        styled += 1
    log(f"Updated SQLHealthcheck static placeholders: {placeholder_updates}; static table fields: {updated}; styled tables: {styled}")


def _is_sql_volume_table(labels: list[str]) -> bool:
    expected = {"volume_mount_point", "file_system_type", "logical_volume_name"}
    return len(expected.intersection(labels)) >= 2


def update_sql_volume_table(table, volume_rows: list[dict[str, str]]) -> int:
    if not volume_rows:
        return 0
    while len(table.columns) < len(volume_rows) + 1:
        table.add_column(Inches(8.8 / 2.54))

    updated = 0
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label = _cell_text(row.cells[0])
        for index, volume in enumerate(volume_rows, start=1):
            if index >= len(row.cells):
                continue
            value = volume.get(label, "")
            if value:
                _set_cell_text(row.cells[index], value)
                updated += 1
    return updated


def replace_sql_static_placeholders(doc: DocumentObject, values: dict[str, str]) -> int:
    placeholders = {
        "{{disk_volume_mount_point}}": values.get("volume_mount_point", ""),
        "{{disk_file_system_type}}": values.get("file_system_type", ""),
        "{{disk_logical_volume_name}}": values.get("logical_volume_name", ""),
        "{{disk_total_size_gb}}": values.get("Total Size (GB)", ""),
        "{{disk_available_size_gb}}": values.get("Available Size (GB)", ""),
        "{{disk_space_free_pct}}": values.get("Space Free %", ""),
        "{{machine_name}}": values.get("Machine Name", ""),
        "{{operating_system}}": values.get("Operating System", ""),
        "{{domain}}": values.get("Domain", ""),
        "{{physical_processors}}": values.get("Physical Processors", ""),
        "{{processor_cores}}": values.get("Processor Cores", ""),
        "{{logical_processors}}": values.get("Logical Processors", ""),
        "{{physical_memory_mb}}": values.get("Physical Memory (MB)", ""),
        "{{instance_name}}": values.get("Instance Name", ""),
        "{{cluster}}": values.get("Cluster", ""),
        "{{sql_server_version}}": values.get("SQL Server Version", ""),
        "{{collation}}": values.get("Collation", ""),
        "{{data_path}}": values.get("Data Path", ""),
        "{{log_path}}": values.get("Log Path", ""),
    }

    updated = 0
    for placeholder, value in placeholders.items():
        if value and replace_placeholder_text(doc, placeholder, value):
            updated += 1
    return updated


def build_sql_volume_values(xls: pd.ExcelFile) -> dict[str, str]:
    volume_rows = build_sql_volume_rows(xls)
    return volume_rows[0] if volume_rows else {}


def build_sql_volume_rows(xls: pd.ExcelFile) -> list[dict[str, str]]:
    if "Volume Info" not in xls.sheet_names:
        return []
    dataframe = pd.read_excel(xls, sheet_name="Volume Info")
    dataframe = dataframe.where(pd.notna(dataframe), "")
    if dataframe.empty:
        return []

    labels = [
        "volume_mount_point",
        "file_system_type",
        "logical_volume_name",
        "Total Size (GB)",
        "Available Size (GB)",
        "Space Free %",
    ]
    rows = []
    for _, row in dataframe.iterrows():
        values = {
            label: _format_static_table_value(row[label])
            for label in labels
            if label in dataframe.columns and _clean_text(row[label])
        }
        if values:
            rows.append(values)
    return rows


def build_sql_hardware_values(xls: pd.ExcelFile) -> dict[str, str]:
    server = _first_sheet_row(xls, "Server Properties")
    hardware = _first_sheet_row(xls, "Hardware Info")
    memory = _first_sheet_row(xls, "System Memory")
    processor = _first_sheet_row(xls, "Processor Description")
    version = _first_sheet_row(xls, "Version Info")

    server_name = _clean_text(server.get("ServerName") or version.get("Server Name") or "")
    instance_name = _clean_text(
        server.get("InstanceName")
        or server.get("Instance Name")
        or (server_name.split("\\", 1)[1] if "\\" in server_name else "")
    )
    sql_version = _sql_version_summary(version.get("SQL Server and OS Version Info") or "")
    return {
        "Instance Name": instance_name,
        "Machine Name": _clean_text(server.get("MachineName") or _machine_from_server_name(server_name)),
        "Operating System": _os_summary(version.get("SQL Server and OS Version Info") or ""),
        "Domain": _clean_text(
            server.get("Domain")
            or server.get("DomainName")
            or server.get("Domain Name")
            or server.get("ServerDomain")
            or server.get("Server Domain")
            or "N/A"
        ),
        "Physical Processors": _processor_model(processor),
        "Processor Cores": _processor_core_count(xls, hardware),
        "Logical Processors": _logical_processor_count(xls, hardware),
        "Physical Memory (MB)": _clean_text(memory.get("Physical Memory (MB)") or hardware.get("Physical Memory (MB)") or ""),
        "Cluster": _cluster_text(server.get("IsClustered") or ""),
        "SQL Server Version": sql_version,
        "Collation": _clean_text(server.get("Collation") or ""),
        "Data Path": _clean_text(server.get("InstanceDefaultDataPath") or ""),
        "Log Path": _clean_text(server.get("InstanceDefaultLogPath") or ""),
    }


def update_sql_summary_table(doc: DocumentObject, xls: pd.ExcelFile, log: LogCallback, language: str = "vi") -> None:
    rows = build_sql_summary_rows(xls, language)
    placeholder_updates = replace_sql_summary_placeholders(doc, rows)

    for table in doc.tables:
        if not table.rows or "HẠNG MỤC" not in _row_text(table.rows[0]):
            continue
        style_sql_summary_table(table)
        if placeholder_updates:
            log(f"Updated SQLHealthcheck summary placeholders: {placeholder_updates}")
            return
        if not rows:
            return
        existing_importance = {
            _cell_text(row.cells[0]): _cell_text(row.cells[4])
            for row in table.rows[1:]
            if len(row.cells) >= 5
        }
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[1]._tr)
        for item in rows:
            row = table.add_row()
            cells = row.cells
            cells[0].text = item["category"]
            cells[1].text = item["assessment"]
            cells[2].text = item["risk"]
            cells[3].text = item["recommendation"]
            cells[4].text = existing_importance.get(item["category"], "")
            for cell in cells:
                format_cell(cell)
            format_cell(cells[0], bold=True)
        log(f"Updated SQLHealthcheck summary rows: {len(rows)}")
        return


def replace_sql_summary_placeholders(doc: DocumentObject, rows: list[dict[str, str]]) -> int:
    key_map = {
        "3.1": "disk",
        "51.": "query",
        "5.2": "missing_index",
        "6 ": "backup",
    }
    placeholders: dict[str, str] = {
        "{{summary_disk_assessment}}": "N/A",
        "{{summary_disk_risk}}": "N/A",
        "{{summary_disk_recommendation}}": "N/A",
        "{{summary_query_assessment}}": "N/A",
        "{{summary_query_risk}}": "N/A",
        "{{summary_query_recommendation}}": "N/A",
        "{{summary_missing_index_assessment}}": "N/A",
        "{{summary_missing_index_risk}}": "N/A",
        "{{summary_missing_index_recommendation}}": "N/A",
        "{{summary_backup_assessment}}": "N/A",
        "{{summary_backup_risk}}": "N/A",
        "{{summary_backup_recommendation}}": "N/A",
    }
    for row in rows:
        prefix = next((value for key, value in key_map.items() if row["category"].startswith(key)), "")
        if not prefix:
            continue
        placeholders[f"{{{{summary_{prefix}_assessment}}}}"] = row["assessment"]
        placeholders[f"{{{{summary_{prefix}_risk}}}}"] = row["risk"]
        placeholders[f"{{{{summary_{prefix}_recommendation}}}}"] = row["recommendation"]

    updated = 0
    for placeholder, value in placeholders.items():
        if value and replace_placeholder_text(doc, placeholder, value):
            updated += 1
    return updated


def build_sql_summary_rows(xls: pd.ExcelFile, language: str = "vi") -> list[dict[str, str]]:
    english = _is_english(language)
    rows: list[dict[str, str]] = []
    disk = _disk_summary(xls, language)
    if disk:
        disk_low_space = _disk_has_low_space(xls, threshold=10)
        rows.append(
            {
                "category": "3.1 Hard disk capacity" if english else "3.1 Dung lượng hard disk",
                "assessment": disk,
                "risk": "May affect storage capacity and operations if disk usage continues to grow." if disk_low_space and english else ("Có thể ảnh hưởng đến khả năng lưu trữ và vận hành nếu dung lượng tiếp tục tăng." if disk_low_space else "N/A"),
                "recommendation": "Monitor capacity growth and expand the disk when it approaches the warning threshold." if disk_low_space and english else ("Theo dõi tăng trưởng dung lượng và mở rộng disk khi gần ngưỡng cảnh báo." if disk_low_space else "N/A"),
            }
        )

    query_count = _sheet_row_count(xls, "Top Worker Time Queries") or _sheet_row_count(xls, "Query Execution Counts")
    if query_count:
        rows.append(
            {
                "category": "5.1 Queries consuming the most resources" if english else "51. Các câu lệnh khi chạy chiếm nhiều tài nguyên",
                "assessment": "Database queries consume significant resources during execution." if english else "Các câu lệnh của các database khi chạy chiếm nhiều tài nguyên khi thực hiện",
                "risk": "Performance impact" if english else "Ảnh hưởng performance",
                "recommendation": "Tune the queries." if english else "Tuning lại câu lệnh",
            }
        )

    missing_index_count = _sheet_row_count(xls, "Missing Indexes")
    if missing_index_count:
        rows.append(
            {
                "category": "5.2 Missing Indexes",
                "assessment": "Some database tables do not have indexes." if english else "Một số table của database không được đánh index",
                "risk": "Performance impact" if english else "Ảnh hưởng đến performance",
                "recommendation": "Create additional indexes for the tables suggested in section 5.2." if english else "Tạo thêm index cho các table được gợi ý trong phần 5.2",
            }
        )

    backup_count, missing_backup_count = _backup_summary_counts(xls)
    if backup_count or missing_backup_count:
        if backup_count and missing_backup_count:
            backup_assessment = (
                f"Backup information was recorded; however, {missing_backup_count} database(s) do not have Last Full Backup recorded."
                if english
                else f"Ghi nhận có thông tin backup, tuy nhiên có {missing_backup_count} database chưa ghi nhận Last Full Backup."
            )
        elif backup_count:
            backup_assessment = "All databases are fully backed up." if english else "Ghi nhận các database backup đầy đủ."
        else:
            backup_assessment = "Databases are not fully backed up." if english else "Các Database không được backup đầy đủ"
        backup_recommendation = (
            "Prepare an environment to perform backup restore testing. Without a restore test environment, backup recoverability cannot be confirmed when needed."
            if english
            else "Khuyến nghị chuẩn bị môi trường thực hiện kiểm thử restore các bản backup. Việc không có môi trường khôi phục kiểm thử bản backup sẽ không đảm bảo bản backup có thể khôi phục thành công khi cần thiết."
        )
        rows.append(
            {
                "category": "6 Backup and recovery" if english else "6 Sao lưu và phục hồi",
                "assessment": backup_assessment,
                "risk": "Data loss when restore is required" if missing_backup_count and english else ("Mất dữ liệu khi cần restore" if missing_backup_count else "N/A"),
                "recommendation": "Perform complete backups for all databases." if missing_backup_count and english else ("Backup đầy đủ cho các databases" if missing_backup_count else backup_recommendation),
            }
        )

    io_warning_count = _sheet_row_count(xls, "IO Warnings")
    if io_warning_count:
        rows.append(
            {
                "category": "3.2 I/O warning",
                "assessment": f"{io_warning_count} I/O warning(s) were recorded." if english else f"Ghi nhận {io_warning_count} cảnh báo I/O.",
                "risk": "May affect read/write performance" if english else "Có thể ảnh hưởng hiệu năng đọc/ghi",
                "recommendation": "Check storage, latency, and databases that generated warnings." if english else "Kiểm tra storage, latency và các database phát sinh cảnh báo",
            }
        )
    return rows


def _first_sheet_row(xls: pd.ExcelFile, sheet_name: str) -> dict[str, str]:
    resolved_sheet = resolve_sql_sheet_name(xls, sheet_name)
    if not resolved_sheet:
        return {}
    dataframe = pd.read_excel(xls, sheet_name=resolved_sheet)
    dataframe = dataframe.where(pd.notna(dataframe), "")
    if dataframe.empty:
        return {}
    row = dataframe.iloc[0]
    return {str(column): _clean_text(row[column]) for column in dataframe.columns}


def _sheet_row_count(xls: pd.ExcelFile, sheet_name: str) -> int:
    resolved_sheet = resolve_sql_sheet_name(xls, sheet_name)
    if not resolved_sheet:
        return 0
    dataframe = pd.read_excel(xls, sheet_name=resolved_sheet)
    return int(len(dataframe.index))


def resolve_sql_sheet_name(xls: pd.ExcelFile, sheet_name: str) -> str | None:
    candidates = SQL_SHEET_ALIASES.get(sheet_name, [sheet_name])
    for candidate in candidates:
        if candidate in xls.sheet_names:
            return candidate
    normalized = {name.strip().lower(): name for name in xls.sheet_names}
    for candidate in candidates:
        found = normalized.get(candidate.strip().lower())
        if found:
            return found
    return None


def _backup_summary_counts(xls: pd.ExcelFile) -> tuple[int, int]:
    backup_count = 0
    backup_names: set[str] = set()
    recent_backup_sheet = resolve_sql_sheet_name(xls, "Recent Full Backups")
    if recent_backup_sheet:
        backups = pd.read_excel(xls, sheet_name=recent_backup_sheet)
        backup_count = int(len(backups.index))
        if "Database Name" in backups.columns:
            backup_names = {_clean_text(value).lower() for value in backups["Database Name"] if _clean_text(value)}

    last_backup_sheet = resolve_sql_sheet_name(xls, "Last Backup By Database")
    if last_backup_sheet:
        last_backups = pd.read_excel(xls, sheet_name=last_backup_sheet)
        last_backups = last_backups.where(pd.notna(last_backups), "")
        full_backup_col = _column_named(last_backups, "Last Full Backup")
        if full_backup_col:
            missing_count = int(sum(1 for value in last_backups[full_backup_col] if not _clean_text(value)))
            return backup_count, missing_count

    return backup_count, 0 if backup_names else 0


def _disk_summary(xls: pd.ExcelFile, language: str = "vi") -> str:
    english = _is_english(language)
    if "Volume Info" not in xls.sheet_names:
        return ""
    dataframe = pd.read_excel(xls, sheet_name="Volume Info")
    dataframe = dataframe.where(pd.notna(dataframe), "")
    if dataframe.empty:
        return ""
    percent_col = _column_named(dataframe, "Space Free %")
    available_col = _column_named(dataframe, "Available Size (GB)")
    mount_col = _column_named(dataframe, "volume_mount_point")
    if not percent_col:
        return ""

    numeric_percent = pd.to_numeric(dataframe[percent_col], errors="coerce")
    if numeric_percent.isna().all():
        return ""

    parts = []
    for index, free_percent in numeric_percent.dropna().items():
        available = _to_number(dataframe.loc[index, available_col]) if available_col else None
        mount = _clean_text(dataframe.loc[index, mount_col]) if mount_col else ""
        location = f" {mount}" if mount and len(numeric_percent.dropna()) > 1 else ""
        if available is not None:
            if english:
                parts.append(f"- Server hard disk{location} currently has {_format_decimal(float(free_percent))}% free (~{_format_decimal(available)}GB)")
            else:
                parts.append(f"- Dung lượng ổ cứng{location} của server hiện tại còn {_format_decimal(float(free_percent))}% trống (~{_format_decimal(available)}GB)")
        else:
            if english:
                parts.append(f"- Server hard disk{location} currently has {_format_decimal(float(free_percent))}% free")
            else:
                parts.append(f"- Dung lượng ổ cứng{location} của server hiện tại còn {_format_decimal(float(free_percent))}% trống")
    return "\n".join(parts)


def _disk_has_low_space(xls: pd.ExcelFile, threshold: float = 10) -> bool:
    if "Volume Info" not in xls.sheet_names:
        return False
    dataframe = pd.read_excel(xls, sheet_name="Volume Info")
    if dataframe.empty:
        return False
    percent_col = _column_named(dataframe, "Space Free %")
    if not percent_col:
        return False
    numeric_percent = pd.to_numeric(dataframe[percent_col], errors="coerce")
    return bool((numeric_percent < threshold).any())


def _normalize_report_language(value: str | None) -> str:
    return "en" if _is_english(value) else "vi"


def _is_english(value: str | None) -> bool:
    return str(value or "").strip().lower() in {"en", "eng", "english"}


def _column_named(dataframe: pd.DataFrame, expected: str) -> str | None:
    expected_norm = expected.strip().lower()
    for column in dataframe.columns:
        if str(column).strip().lower() == expected_norm:
            return column
    return None


def _clean_text(value) -> str:
    if value is None or pd.isna(value):
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def _to_number(value) -> float | None:
    text = _clean_text(value).replace(",", "")
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    return float(match.group(0)) if match else None


def _format_decimal(value: float) -> str:
    return f"{value:.2f}".rstrip("0").rstrip(".")


def _format_static_table_value(value) -> str:
    if value is None or pd.isna(value):
        return ""
    numeric_value = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
    if pd.notna(numeric_value):
        return _format_decimal(float(numeric_value))
    return _clean_text(value)


def _machine_from_server_name(server_name: str) -> str:
    return server_name.split("\\", 1)[0].strip()


def _sql_version_summary(raw_version: str) -> str:
    text = _clean_text(raw_version)
    match = re.search(r"(Microsoft SQL Server\s+\d{4})", text, re.IGNORECASE)
    return match.group(1) if match else text


def _os_summary(raw_version: str) -> str:
    text = _clean_text(raw_version)
    match = re.search(r"on\s+(.+?)(?:\s+\(Build|\s*$)", text, re.IGNORECASE)
    if not match:
        return ""
    return match.group(1).replace("Windows NT 6.3", "Windows Server 2012 R2").strip()


def _processor_model(processor: dict[str, str]) -> str:
    model = _clean_text(processor.get("Data") or processor.get("Processor Name") or processor.get("ProcessorNameString") or "")
    return re.sub(r"-\s+", "-", model)


def _processor_core_count(xls: pd.ExcelFile, hardware: dict[str, str]) -> str:
    for field in [
        "Total Physical Cores",
        "Total Physical Core Count",
        "Physical Core Count",
        "CPU Core Count",
        "Core Count",
    ]:
        value = _clean_text(hardware.get(field) or "")
        if value:
            return _format_static_table_value(value)

    core_count = _cores_from_log(xls)
    if core_count:
        return core_count

    for field in ["Processor Core Count", "Cores"]:
        value = _clean_text(hardware.get(field) or "")
        if value:
            return _format_static_table_value(value)
    return ""


def _logical_processor_count(xls: pd.ExcelFile, hardware: dict[str, str]) -> str:
    for field in ["Logical CPU Count", "Logical Processor Count", "Total Logical Processors"]:
        value = _clean_text(hardware.get(field) or "")
        if value:
            return _format_static_table_value(value)

    logical_count = _logical_processors_from_log(xls)
    return logical_count or ""


def _cores_from_log(xls: pd.ExcelFile) -> str:
    if "Core Counts" not in xls.sheet_names:
        return ""
    dataframe = pd.read_excel(xls, sheet_name="Core Counts")
    dataframe = dataframe.where(pd.notna(dataframe), "")
    if dataframe.empty:
        return ""

    row = dataframe.iloc[0]
    total_column_value = _first_numeric_row_value(
        row,
        [
            "Total Physical Cores",
            "Total Physical Core Count",
            "Physical Core Count",
            "Processor Cores",
        ],
    )
    if total_column_value:
        return total_column_value

    socket_count = _first_numeric_row_value(row, ["Socket Count", "Sockets", "socket_count"])
    cores_per_socket = _first_numeric_row_value(row, ["Cores Per Socket", "cores_per_socket"])
    if socket_count and cores_per_socket:
        return str(int(float(socket_count)) * int(float(cores_per_socket)))

    if "Text" not in dataframe.columns:
        return ""
    text = _clean_text(dataframe.iloc[0]["Text"])
    total_match = re.search(r"(\d+)\s+total\s+physical\s+cores?", text, re.IGNORECASE)
    if total_match:
        return total_match.group(1)
    socket_match = re.search(r"(\d+)\s+sockets?\s+with\s+(\d+)\s+cores?\s+per\s+socket", text, re.IGNORECASE)
    if socket_match:
        return str(int(socket_match.group(1)) * int(socket_match.group(2)))
    socket_count_match = re.search(r"socket\s+count\s*[:=]\s*(\d+)", text, re.IGNORECASE)
    cores_per_socket_match = re.search(r"cores?\s+per\s+socket\s*[:=]\s*(\d+)", text, re.IGNORECASE)
    if socket_count_match and cores_per_socket_match:
        return str(int(socket_count_match.group(1)) * int(cores_per_socket_match.group(1)))
    return ""


def _logical_processors_from_log(xls: pd.ExcelFile) -> str:
    if "Core Counts" not in xls.sheet_names:
        return ""
    dataframe = pd.read_excel(xls, sheet_name="Core Counts")
    dataframe = dataframe.where(pd.notna(dataframe), "")
    if dataframe.empty:
        return ""
    row = dataframe.iloc[0]
    total_column_value = _first_numeric_row_value(row, ["Logical CPU Count", "Logical Processor Count", "Total Logical Processors"])
    if total_column_value:
        return total_column_value
    if "Text" not in dataframe.columns:
        return ""
    text = _clean_text(dataframe.iloc[0]["Text"])
    total_match = re.search(r"(\d+)\s+total\s+logical\s+processors?", text, re.IGNORECASE)
    return total_match.group(1) if total_match else ""


def _first_numeric_row_value(row: pd.Series, fields: list[str]) -> str:
    normalized = {str(column).strip().lower(): column for column in row.index}
    for field in fields:
        column = normalized.get(field.strip().lower())
        if column is None:
            continue
        number = _to_number(row[column])
        if number is not None:
            return _format_decimal(number)
    return ""


def _cluster_text(value) -> str:
    text = _clean_text(value).lower()
    if text in {"0", "false", "no"}:
        return "None"
    if text in {"1", "true", "yes"}:
        return "Clustered"
    return _clean_text(value)


def _cell_text(cell: _Cell) -> str:
    return re.sub(r"\s+", " ", cell.text).strip()


def _row_text(row) -> str:
    return " | ".join(_cell_text(cell) for cell in row.cells)


def _set_cell_text(cell: _Cell, value: str) -> None:
    cell.text = value
    format_cell(cell)


def _set_cell_text_preserve_format(cell: _Cell, value: str) -> None:
    if not cell.paragraphs:
        cell.text = value
        return
    _set_paragraph_text(cell.paragraphs[0], value)
    for paragraph in cell.paragraphs[1:]:
        _set_paragraph_text(paragraph, "")


def style_sql_key_value_table(table) -> None:
    if not table.rows or len(table.rows[0].cells) < 2:
        return
    table.autofit = False
    set_table_layout(table, fixed=True)
    set_table_width_percent(table, 100)
    set_table_borders(table)
    if len(table.columns) > 2:
        set_column_width(table.columns[0], 5.0)
        data_width = max(3.8, min(8.8, 11.4 / (len(table.columns) - 1)))
        for column_index in range(1, len(table.columns)):
            set_column_width(table.columns[column_index], data_width)
    else:
        set_column_width(table.columns[0], 7.6)
        set_column_width(table.columns[1], 8.8)
    for row in table.rows:
        set_row_height(row, 0.65, rule="atLeast")
        for column_index, cell in enumerate(row.cells):
            set_cell_wrap(cell)
            format_cell(cell, bold=column_index == 0, font_size=12)


def style_sql_summary_table(table) -> None:
    set_table_width_percent(table, 100)
    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        set_row_height(row, 0.8, rule="atLeast")
        for column_index, cell in enumerate(row.cells):
            set_cell_wrap(cell)
            format_cell(cell, bold=row_index == 0 or column_index == 0, font_size=12)


def estimate_vertical_row_height(row, columns, horizontal_columns: list[str], config: dict) -> float:
    vertical_texts = [
        _clean_text(value)
        for column, value in zip(columns, row)
        if column not in horizontal_columns
    ]
    longest = max((len(value) for value in vertical_texts), default=0)
    estimated = 0.55 + (longest * 0.18)
    min_height = float(config.get("min_row_height", 1.4))
    max_height = float(config.get("max_row_height", 3.6))
    return max(min_height, min(max_height, estimated))


def _replace_value_after_label(doc: DocumentObject, label: str, value: str) -> bool:
    paragraphs = list(iter_all_paragraphs(doc))
    for index, paragraph in enumerate(paragraphs[:-1]):
        if paragraph.text.strip() == label:
            _set_paragraph_text(paragraphs[index + 1], value)
            return True
    return False


def _set_paragraph_text(paragraph: Paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = value


def extract_scalar_value(xls: pd.ExcelFile, config: dict) -> str:
    sheet_name = resolve_sql_sheet_name(xls, config["sheet"])
    if not sheet_name:
        raise ValueError(f"sheet not found '{config['sheet']}'")

    dataframe = pd.read_excel(xls, sheet_name=sheet_name)
    if dataframe.empty:
        return ""

    value_column = config["value_column"]
    selector_column = config.get("selector_column")
    selected_row = dataframe.iloc[0]

    if selector_column and selector_column in dataframe.columns:
        numeric_selector = pd.to_numeric(dataframe[selector_column], errors="coerce")
        if numeric_selector.notna().any():
            if config.get("selector") == "max":
                selected_row = dataframe.loc[numeric_selector.idxmax()]
            else:
                selected_row = dataframe.loc[numeric_selector.idxmin()]

    value = selected_row[value_column] if value_column in dataframe.columns else ""
    numeric_value = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
    if pd.notna(numeric_value) and config.get("format"):
        return config["format"].format(float(numeric_value))
    return "" if pd.isna(value) else str(value)


def _prepare_dataframe(dataframe: pd.DataFrame, config: dict) -> pd.DataFrame:
    dataframe = dataframe.where(pd.notna(dataframe), "")

    if config.get("transpose", False):
        if config.get("columns"):
            selected_cols = [dataframe.columns[i] for i in config["columns"] if i < len(dataframe.columns)]
            dataframe = dataframe[selected_cols]

        original_first_col = dataframe.columns[0]
        dataframe = dataframe.T.reset_index()

        if len(dataframe.columns) > 1:
            new_columns = [original_first_col] + [str(value) for value in dataframe.iloc[0, 1:].tolist()]
            dataframe.columns = new_columns
            dataframe = dataframe.iloc[1:].reset_index(drop=True)

        dataframe = dataframe.loc[:, ~dataframe.columns.astype(str).str.lower().str.contains("nan", na=False)]
        dataframe = dataframe.loc[:, dataframe.columns.astype(str).str.strip() != ""]
        return dataframe

    if config.get("columns"):
        selected = [dataframe.columns[i] for i in config["columns"] if i < len(dataframe.columns)]
        dataframe = dataframe[selected]

    return dataframe


def replace_placeholder_text(doc: DocumentObject, placeholder: str, replacement: str) -> bool:
    replaced = False
    for paragraph in iter_all_paragraphs(doc):
        updated_text = replace_placeholder_in_text(paragraph.text, placeholder, replacement)
        if updated_text != paragraph.text:
            _set_paragraph_text(paragraph, updated_text)
            replaced = True
    return replaced


def replace_placeholder_in_text(text: str, placeholder: str, replacement: str) -> str:
    if placeholder in text:
        return text.replace(placeholder, replacement)
    match = re.fullmatch(r"\{\{([A-Za-z0-9_]+)\}\}", placeholder)
    if not match:
        return text
    pattern = r"\{\{\s*" + re.escape(match.group(1)) + r"\s*\}\}"
    return re.sub(pattern, lambda _match: replacement, text)


def replace_placeholder_with_table(doc: DocumentObject, placeholder: str, dataframe: pd.DataFrame, config: dict) -> bool:
    for paragraph in iter_all_paragraphs(doc):
        if placeholder not in paragraph.text:
            continue

        if config.get("key_value_table"):
            table = doc.add_table(rows=len(dataframe.index), cols=len(dataframe.columns))
            table.autofit = False
            set_table_layout(table, fixed=True)
            set_table_width_percent(table, 100)
            set_table_borders(table)
            if config.get("column_widths"):
                for column_index, width_cm in enumerate(config["column_widths"]):
                    if column_index < len(table.columns):
                        set_column_width(table.columns[column_index], width_cm)
            for row_index, (_, row) in enumerate(dataframe.iterrows()):
                table_row = table.rows[row_index]
                set_row_height(table_row, config.get("row_height", 0.65), rule=config.get("row_height_rule", "atLeast"))
                for column_index, value in enumerate(row):
                    cell = table_row.cells[column_index]
                    cell.text = str(value)
                    set_cell_wrap(cell)
                    format_cell(cell, bold=column_index == 0, font_size=config.get("font_size", 12))
            paragraph._element.getparent().replace(paragraph._element, table._element)
            return True

        table = doc.add_table(rows=1, cols=len(dataframe.columns))
        table.autofit = not bool(config.get("column_widths"))
        set_table_borders(table)
        set_table_layout(table, fixed=bool(config.get("column_widths")))
        set_table_width_percent(table, 100)

        header_height = config.get("header_height", 0.8)
        if header_height:
            set_row_height(table.rows[0], header_height, rule=config.get("header_height_rule", "atLeast"))

        use_vertical_header = config.get("vertical_header", False)
        font_size = config.get("font_size", 12)
        for column_index, column_name in enumerate(dataframe.columns):
            cell = table.rows[0].cells[column_index]
            cell.text = str(column_name)
            set_cell_bg(cell, "0066CC")
            set_cell_wrap(cell)
            format_cell(cell, bold=True, font_color=RGBColor(255, 255, 255), font_size=font_size)
            if use_vertical_header:
                set_cell_text_direction(cell, "tbRl")

        horizontal_columns = config.get("horizontal_columns", [])
        vertical_body = config.get("vertical_body", False)
        row_height = config.get("row_height")

        for _, row in dataframe.iterrows():
            new_row = table.add_row()
            dynamic_row_height = estimate_vertical_row_height(row, dataframe.columns, horizontal_columns, config) if config.get("dynamic_vertical_row_height") else None
            if row_height or dynamic_row_height:
                set_row_height(new_row, row_height or dynamic_row_height, rule=config.get("row_height_rule", "atLeast"))
            for column_index, value in enumerate(row):
                cell = new_row.cells[column_index]
                cell.text = str(value)
                set_cell_wrap(cell)
                format_cell(cell, font_size=font_size)

                if vertical_body and dataframe.columns[column_index] not in horizontal_columns:
                    set_cell_text_direction(cell, "tbRl")

        if config.get("column_widths"):
            for column_index, width_cm in enumerate(config["column_widths"]):
                if column_index < len(table.columns):
                    set_column_width(table.columns[column_index], width_cm)

        paragraph._element.getparent().replace(paragraph._element, table._element)
        return True
    return False


def replace_placeholder_with_image(doc: DocumentObject, placeholder: str, image_path: Path) -> bool:
    for paragraph in iter_all_paragraphs(doc):
        if placeholder not in paragraph.text:
            continue

        paragraph.text = paragraph.text.replace(placeholder, "")
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(5.5))
        return True
    return False


def iter_all_paragraphs(doc: DocumentObject) -> Iterable[Paragraph]:
    yield from _iter_paragraphs(doc)
    for section in doc.sections:
        yield from _iter_paragraphs(section.header)
        yield from _iter_paragraphs(section.footer)


def _iter_paragraphs(parent) -> Iterable[Paragraph]:
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def _replace_text_in_paragraph(paragraph: Paragraph, placeholder: str, replacement: str) -> None:
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)

    if placeholder in paragraph.text:
        paragraph.text = paragraph.text.replace(placeholder, replacement)


def create_pie_chart(
    dataframe: pd.DataFrame,
    title: str,
    output_image: Path,
    label_col_idx: int = 0,
    value_col_idx: int = 1,
    top_n: int = 10,
    log: LogCallback | None = None,
) -> bool:
    log = log or print
    try:
        df_chart = dataframe.head(top_n).copy()
        labels = df_chart.iloc[:, label_col_idx].astype(str).tolist()
        values = pd.to_numeric(df_chart.iloc[:, value_col_idx], errors="coerce").fillna(0).tolist()

        valid_data = [
            (label, value)
            for label, value in zip(labels, values, strict=False)
            if value > 0 and label.strip() != "" and label.lower() not in ["nan", "none"]
        ]
        if not valid_data:
            log(f"No valid data for chart: {title}")
            return False

        labels, values = zip(*valid_data, strict=False)
        fig, ax = plt.subplots(figsize=(10, 8), facecolor="white")
        colors = [
            "#5B9BD5",
            "#ED7D31",
            "#A5A5A5",
            "#FFC000",
            "#70AD47",
            "#4472C4",
            "#C55A11",
            "#7030A0",
            "#44546A",
            "#264478",
        ]
        wedges, _texts = ax.pie(
            values,
            labels=None,
            startangle=90,
            colors=colors[: len(values)],
        )
        ax.set_title(title, fontsize=16, fontweight="bold", pad=30, color="#333333")

        num_cols = 3 if len(labels) > 6 else (2 if len(labels) > 3 else 1)
        legend = ax.legend(labels, loc="upper center", bbox_to_anchor=(0.5, -0.05), ncol=num_cols, frameon=False, fontsize=11)

        for index, _wedge in enumerate(wedges):
            if index < len(legend.get_patches()):
                legend.get_patches()[index].set_facecolor(colors[index % len(colors)])

        plt.axis("equal")
        plt.tight_layout()
        plt.savefig(output_image, dpi=150, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        log(f"Created chart: {output_image}")
        return True
    except Exception as exc:
        log(f"Error creating chart: {exc}")
        plt.close()
        return False


def set_cell_bg(cell: _Cell, fill_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_color)
    tc_pr.append(shading)


def set_cell_text_direction(cell: _Cell, direction: str = "lrTb") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    text_direction = OxmlElement("w:textDirection")
    text_direction.set(qn("w:val"), direction)
    tc_pr.append(text_direction)


def set_cell_wrap(cell: _Cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for no_wrap in tc_pr.findall(qn("w:noWrap")):
        tc_pr.remove(no_wrap)
    set_cell_margins(cell)


def set_cell_margins(cell: _Cell, margin_dxa: int = 60) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ["top", "left", "bottom", "right"]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_dxa))
        node.set(qn("w:type"), "dxa")


def set_row_height(row, height_cm: float, rule: str = "exact") -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for existing_height in tr_pr.findall(qn("w:trHeight")):
        tr_pr.remove(existing_height)
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(height_cm * 567)))
    tr_height.set(qn("w:hRule"), rule)
    tr_pr.append(tr_height)


def format_cell(cell: _Cell, bold: bool = False, font_color: RGBColor | None = None, font_size: int = 12) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        for run in paragraph.runs:
            run.font.name = "Cambria"
            run.font.size = Pt(font_size)
            run.bold = bold
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria")
            if font_color:
                run.font.color.rgb = font_color


def set_table_layout(table, fixed: bool = False) -> None:
    table_properties = table._element.tblPr
    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed" if fixed else "autofit")


def set_table_width_percent(table, percent: int = 100) -> None:
    table_properties = table._element.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(percent * 50))
    table_width.set(qn("w:type"), "pct")


def set_table_borders(table) -> None:
    table_properties = table._element.tblPr
    existing_borders = table_properties.find(qn("w:tblBorders"))
    if existing_borders is not None:
        table_properties.remove(existing_borders)
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    table_properties.append(borders)


def set_column_width(column, width_cm: float) -> None:
    for cell in column.cells:
        cell.width = Inches(width_cm / 2.54)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.first_child_found_in("w:tcW")
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(int((width_cm / 2.54) * 1440)))
        tc_w.set(qn("w:type"), "dxa")


def _make_logger(log_callback: LogCallback | None) -> LogCallback:
    def log(message: str) -> None:
        print(message)
        if log_callback:
            log_callback(message)

    return log
