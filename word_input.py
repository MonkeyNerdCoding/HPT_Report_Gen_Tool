from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def replace_placeholder_with_table(doc, placeholder, table_data):
    """
    Replace placeholder text in Word doc with formatted table.
    """
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = ""
            rows = len(table_data)
            cols = max(len(r) for r in table_data)

            table = doc.add_table(rows=rows, cols=cols)
            table.style = "Table Grid"

            # Format table
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    cell = table.cell(i, j)
                    cell.text = cell_text

                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.name = "Cambria"
                            run.font.size = Pt(12)
                            run.font.color.rgb = RGBColor(0, 0, 0)

                    # Header style
                    if i == 0:
                        shading = parse_xml(
                            r'<w:shd {} w:fill="0066CC"/>'.format(nsdecls("w"))
                        )
                        cell._tc.get_or_add_tcPr().append(shading)
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)

            break  # stop after first match

def save_word(doc, output_path):
    doc.save(output_path)


# New implementation compatibility exports.
from pathlib import Path

from models import MappingRule, TableContent
from rendering.word_renderer import replace_placeholder as _replace_placeholder


def replace_placeholder_with_table(doc, placeholder, table_data):
    """
    Backward-compatible wrapper for the original API.

    New code should call rendering.word_renderer.render_report().
    """
    content = TableContent(source_path=Path("."), rows=table_data, logical_key=placeholder)
    rule = MappingRule(placeholder=placeholder, width_inches=None)
    return _replace_placeholder(doc, placeholder, rule, content)
